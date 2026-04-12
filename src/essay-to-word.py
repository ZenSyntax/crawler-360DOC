"""随笔清洗 HTML 转 Word（.docx）。由 wc-essay 动态加载；主程序为仓库根或 src 下的 wc-essay.py。"""

from __future__ import annotations

import re
from collections.abc import Callable
from pathlib import Path

from bs4 import BeautifulSoup
from bs4.element import NavigableString, Tag
from docx import Document
from docx.enum.text import WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import Pt

# 可选依赖 lxml 作为 BS4 解析器；未安装则回退 html.parser。
try:
    import lxml  # noqa: F401

    _BS_PARSER = "lxml"
except ImportError:
    _BS_PARSER = "html.parser"

FONT_PT_WUHAO = Pt(10.5)
LINE_SPACING_FIXED_PT = Pt(20)


def _beautiful_soup(markup: str) -> BeautifulSoup:
    return BeautifulSoup(markup, _BS_PARSER)


def set_document_defaults(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal.font.size = FONT_PT_WUHAO
    r_pr = normal._element.rPr
    if r_pr is not None and r_pr.rFonts is not None:
        r_pr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    npf = normal.paragraph_format
    npf.space_before = Pt(0)
    npf.space_after = Pt(0)
    npf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    npf.line_spacing = LINE_SPACING_FIXED_PT


def apply_fixed_line_spacing_20pt(paragraph) -> None:
    pf = paragraph.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    pf.line_spacing = LINE_SPACING_FIXED_PT


def normalize_essay_newlines(text: str) -> str:
    text = re.sub(r"\n{2,}", "\n", text)
    return text.strip()


def add_inline_to_paragraph(paragraph, node) -> None:
    if isinstance(node, NavigableString):
        run = paragraph.add_run(str(node))
        run.font.size = FONT_PT_WUHAO
        return
    if not isinstance(node, Tag):
        return
    if node.name in ("strong", "b"):
        run = paragraph.add_run(node.get_text())
        run.bold = True
        run.font.size = FONT_PT_WUHAO
        return
    if node.name == "br":
        paragraph.add_run().add_break()
        return
    for child in node.children:
        add_inline_to_paragraph(paragraph, child)


def add_page_meta(doc: Document, header: Tag | None) -> None:
    if header is None:
        return
    p = doc.add_paragraph()
    for ch in header.children:
        add_inline_to_paragraph(p, ch)
    apply_fixed_line_spacing_20pt(p)


def add_essay_body_paragraph(doc: Document, body_el: Tag | None) -> None:
    if body_el is None:
        return
    text = normalize_essay_newlines(body_el.get_text())
    p = doc.add_paragraph()
    apply_fixed_line_spacing_20pt(p)
    lines = text.split("\n")
    for i, line in enumerate(lines):
        if i > 0:
            p.add_run().add_break()
        run = p.add_run(line)
        run.font.size = FONT_PT_WUHAO


def add_essay_article(doc: Document, article: Tag) -> None:
    date_el = article.select_one(".essay-date")
    if date_el is not None:
        p = doc.add_paragraph()
        run = p.add_run(date_el.get_text(strip=True))
        run.bold = True
        run.font.size = FONT_PT_WUHAO
        apply_fixed_line_spacing_20pt(p)

    body_el = article.select_one(".essay-body")
    add_essay_body_paragraph(doc, body_el)


def html_file_to_docx(html_path: Path, docx_path: Path, *, force: bool) -> bool:
    if docx_path.exists() and not force:
        return False
    raw = html_path.read_text(encoding="utf-8", errors="replace")
    soup = _beautiful_soup(raw)

    doc = Document()
    set_document_defaults(doc)

    add_page_meta(doc, soup.select_one("header.page-meta"))

    main = soup.select_one("main")
    if main is not None:
        for nd in main.select("p.no-data"):
            p = doc.add_paragraph(normalize_essay_newlines(nd.get_text()))
            for r in p.runs:
                r.font.size = FONT_PT_WUHAO
            apply_fixed_line_spacing_20pt(p)
        for art in main.select("article.essay"):
            if isinstance(art, Tag):
                add_essay_article(doc, art)
    else:
        body = soup.body
        if body is not None:
            p = doc.add_paragraph(
                normalize_essay_newlines(body.get_text("\n", strip=True))
            )
            for r in p.runs:
                r.font.size = FONT_PT_WUHAO
            apply_fixed_line_spacing_20pt(p)

    doc.save(str(docx_path))
    return True


def _docx_needs_regen(
    html_path: Path, docx_path: Path, *, force: bool, incremental: bool
) -> bool:
    if force:
        return True
    if not docx_path.exists():
        return True
    if not incremental:
        return False
    try:
        return html_path.stat().st_mtime > docx_path.stat().st_mtime
    except OSError:
        return True


def convert_essay_html_tree_to_docx(
    root: Path,
    *,
    force: bool,
    incremental_docx: bool = True,
    log_info: Callable[[str], None],
    log_warn: Callable[[str], None],
) -> int:
    # 递归扫描 root 下 .html 转同名 .docx；返回失败数。
    if not root.is_dir():
        log_warn(f"Word 转换跳过：目录不存在 {root}")
        return 0

    html_files = sorted(p for p in root.rglob("*.html") if p.is_file())
    if not html_files:
        log_warn(f"Word 转换：未找到 HTML under {root}")
        return 0

    incremental = incremental_docx and not force
    ok = 0
    skipped = 0
    failed = 0
    for html_path in html_files:
        docx_path = html_path.with_suffix(".docx")
        try:
            if not _docx_needs_regen(
                html_path, docx_path, force=force, incremental=incremental_docx
            ):
                skipped += 1
                continue
            if html_file_to_docx(html_path, docx_path, force=True):
                log_info(f"[docx] {docx_path}")
                ok += 1
            else:
                skipped += 1
        except Exception as exc:
            log_warn(f"[docx] 失败 {html_path}: {exc}")
            failed += 1

    mode = "增量(mtime)" if incremental else ("强制全量" if force else "仅补缺")
    log_info(
        f"Word 转换完成 [{mode}]: 新生成 {ok}，跳过 {skipped}，失败 {failed}，"
        f"共扫描 {len(html_files)} 个 .html"
    )
    return failed


if __name__ == "__main__":
    raise SystemExit(
        "入口：仓库根 wc-essay.py 或 src/wc-essay.py（本文件为转换实现，不作为主程序）。"
    )
