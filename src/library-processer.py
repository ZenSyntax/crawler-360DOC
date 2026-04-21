"""
Library cleaner module: transforms raw HTML into clean_ HTML and optional Word (.docx).

Loaded by wc-library; entry scripts are wc-library.py under repo root or src/.
Resource directory and clean_ HTML main file share the same stem (without .html).
"""

from __future__ import annotations

import copy
import concurrent.futures
import hashlib
import io
import json
import os
import random
import re
import shutil
import tempfile
import threading
import time
from collections.abc import Callable
from pathlib import Path
from typing import NamedTuple
from urllib.parse import parse_qsl, unquote, urlencode, urljoin, urlparse, urlunparse

import requests
from bs4 import BeautifulSoup
from bs4.element import NavigableString, Tag
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_UNDERLINE
from docx.image.exceptions import UnrecognizedImageError
from docx.text.paragraph import Paragraph
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.oxml import OxmlElement
from docx.oxml.ns import nsdecls, qn
from docx.oxml.parser import parse_xml
from docx.shared import Inches, Pt, RGBColor
from PIL import Image

BASE_URL = "http://www.360doc.com"
IMG_CHANGE_URL = f"{BASE_URL}/Ajax/imgurl.ashx?op=changeurl"
ALLOWED_HOSTS = ("360doc.com", "360doc.cn")
TIMEOUT = 20
MAX_RETRY = 2
RESOURCE_REQUEST_TIMEOUT = 8
RESOURCE_REQUEST_RETRIES = 1
RESOURCE_REQUEST_SLEEP_SEC = (0.2, 0.55)
AFTER_ARTICLE_WITH_RESOURCES_SLEEP_SEC = (0.35, 0.9)
RESOURCE_START_JITTER_SEC = (0.01, 0.08)
# 资源并发线程上限：按资源数动态创建线程，每个线程尽量只处理一个资源。
def _get_env_positive_int(name: str, default: int) -> int:
    raw = os.getenv(name, "").strip()
    if not raw:
        return default
    try:
        val = int(raw)
    except Exception:
        return default
    return val if val > 0 else default


def _get_env_positive_float(name: str, default: float) -> float:
    raw = os.getenv(name, "").strip()
    if not raw:
        return default
    try:
        val = float(raw)
    except Exception:
        return default
    return val if val > 0 else default


RESOURCE_DOWNLOAD_MAX_WORKERS = _get_env_positive_int("DOC360_MAX_WORKERS", 50)
RESOURCE_MAX_ATTEMPTS_PER_URL = 12
RESOURCE_MAX_REFRESH_RETRIES = 4
RESOURCE_PROGRESS_HEARTBEAT_SEC = 8
INVALID_NAME_RE = re.compile(r'[<>:"/\\|?*\x00-\x1F]')
DATE_RE = re.compile(r"\d{4}-\d{2}-\d{2}")
WORD_META_WORDURL_RE = re.compile(r"wordurl\s*=\s*['\"]([^'\"]+)['\"]", re.I)
WORD_META_WORDCSSURL_RE = re.compile(r"wordCSSUrl\s*=\s*['\"]([^'\"]+)['\"]", re.I)
WORD_META_PAGENUM_RE = re.compile(r"pageNume\s*=\s*(\d+)", re.I)
PPT_IMG_ARR_RE = re.compile(r"var\s+pptimgArr\s*=\s*\[(.*?)\]\s*;?", re.I | re.S)
PDF_IMG_ARR_RE = re.compile(r"var\s+pdfList\s*=\s*\[(.*?)\]\s*;?", re.I | re.S)
PPT_IMG_URL_RE = re.compile(r"['\"]([^'\"]+)['\"]")
GER_LOOKING_USER_INFO_RE = re.compile(
    r"GerLookingUserInfo\s*\((.*?)\)\s*;?", re.I | re.S
)
PDF_FIRST_PAGE_URL_RE = re.compile(
    r"['\"]([^'\"]+?_(\d+)\.(?:jpe?g|png|gif|webp|bmp))['\"]", re.I
)
CSS_URL_FUNC_RE = re.compile(r"url\(\s*(['\"]?)([^'\"\)]+)\1\s*\)", re.I)
HTML_CHARSET_META_RE = re.compile(
    rb"<meta[^>]+charset\s*=\s*['\"]?\s*([a-zA-Z0-9._-]+)", re.I
)
WORD_PREVIEW_PAGE_SLEEP_SEC = (0.25, 0.65)
WORD_PREVIEW_REQUEST_TIMEOUT = _get_env_positive_float(
    "DOC360_WORD_PREVIEW_TIMEOUT_SEC", float(RESOURCE_REQUEST_TIMEOUT)
)
WORD_PREVIEW_REQUEST_RETRIES = _get_env_positive_int(
    "DOC360_WORD_PREVIEW_RETRIES", RESOURCE_REQUEST_RETRIES + 1
) - 1
WORD_PREVIEW_PROGRESS_EVERY_PAGES = _get_env_positive_int(
    "DOC360_WORD_PREVIEW_PROGRESS_EVERY", 10
)
WORD_PREVIEW_HEARTBEAT_SEC = _get_env_positive_float(
    "DOC360_WORD_PREVIEW_HEARTBEAT_SEC", 15.0
)
WORD_PREVIEW_MAX_CONSECUTIVE_FAILURES = _get_env_positive_int(
    "DOC360_WORD_PREVIEW_MAX_CONSEC_FAILS", 40
)
MAX_WORD_PREVIEW_PAGES = _get_env_positive_int("DOC360_MAX_WORD_PREVIEW_PAGES", 2000)

# clean_ prefix marks cleaned HTML and differentiates it from raw files during scans.
CLEAN_HTML_PREFIX = "clean_"

# Font-size mapping: title ~15pt, metadata/body ~10.5pt.
TITLE_PT = Pt(15)
META_PT = Pt(10.5)
DEFAULT_BODY_PT = Pt(10.5)
# Unified spacing: fixed 20pt line spacing, zero paragraph spacing.
FIXED_LINE_SPACING_PT = Pt(20)

# Local image suffixes used for src/href checks and src fallback.
_LOCAL_IMAGE_HREF_EXTS = (
    ".jpg",
    ".jpeg",
    ".png",
    ".gif",
    ".webp",
    ".bmp",
)
# Expand wrapper roots during cleaning to avoid adding an extra container div.
_CONTENT_WRAPPER_IDS = frozenset({"artContent", "printArticle"})

CLEAN_ERROR_URL_FILE = Path("logs/clean_error_url.txt")
CLEAN_ARTICLE_ERROR_FILE = Path("logs/clean_article_error.txt")
RESOURCES_NOT_FOUND_WARNING_FILE = Path("logs/resources_not_found_warning.txt")
RATE_LIMIT_STATUS_CODES = {403}
TRANSIENT_GATEWAY_STATUS_CODES = {502, 503, 504}
TRANSIENT_GATEWAY_RETRY_SLEEP_SEC = (0.6, 1.6)
_CATEGORY_ARTNUM_BY_DIR_ID: dict[str, str] = {}

_log_info = print
_log_warn = print


class CleanRateLimitError(RuntimeError):
    pass


class ResourceNotFoundError(RuntimeError):
    pass


class ResourceExpiredError(RuntimeError):
    pass


class ResourceGatewayError(RuntimeError):
    pass


class ResourceLocalizationResult(NamedTuple):
    downloaded: int
    failed_urls: list[str]


def set_processer_loggers(log_info_fn, log_warn_fn) -> None:
    global _log_info, _log_warn
    _log_info = log_info_fn
    _log_warn = log_warn_fn


def set_clean_error_url_file(path: Path) -> None:
    global CLEAN_ERROR_URL_FILE
    CLEAN_ERROR_URL_FILE = path


def set_clean_article_error_file(path: Path) -> None:
    global CLEAN_ARTICLE_ERROR_FILE
    CLEAN_ARTICLE_ERROR_FILE = path


def set_resources_not_found_warning_file(path: Path) -> None:
    global RESOURCES_NOT_FOUND_WARNING_FILE
    RESOURCES_NOT_FOUND_WARNING_FILE = path


def set_category_artnum_map(mapping: dict[str, str]) -> None:
    global _CATEGORY_ARTNUM_BY_DIR_ID
    _CATEGORY_ARTNUM_BY_DIR_ID = {
        str(k).strip(): str(v).strip()
        for k, v in (mapping or {}).items()
        if str(k).strip() and str(v).strip()
    }


def log_info(msg: str) -> None:
    _log_info(msg)


def log_warn(msg: str) -> None:
    _log_warn(msg)


def append_clean_error_url_line(line: str) -> None:
    try:
        CLEAN_ERROR_URL_FILE.parent.mkdir(parents=True, exist_ok=True)
        with CLEAN_ERROR_URL_FILE.open("a", encoding="utf-8") as fp:
            fp.write(f"{line}\n")
    except Exception as exc:
        log_warn(f"写入 clean_error_url.txt 失败 line={line!r} err={exc}")


def append_clean_article_error_line(line: str) -> None:
    try:
        CLEAN_ARTICLE_ERROR_FILE.parent.mkdir(parents=True, exist_ok=True)
        with CLEAN_ARTICLE_ERROR_FILE.open("a", encoding="utf-8") as fp:
            fp.write(f"{line}\n")
    except Exception as exc:
        log_warn(f"写入 clean_article_error.txt 失败 line={line!r} err={exc}")


def append_clean_resource_failure_line(
    *,
    article_id: str,
    article_title: str,
    article_dir_name: str,
    resource_url: str,
    error: Exception,
) -> None:
    # 兼容旧日志前缀：article_id-url-exc；追加 article/dir 便于定位。
    legacy = f"{article_id}-{resource_url}-{error}"
    append_clean_error_url_line(
        f"{legacy}\tarticle={article_title or 'unknown'}\tdir={article_dir_name or 'unknown'}"
    )


def append_resource_not_found_warning_line(line: str) -> None:
    try:
        RESOURCES_NOT_FOUND_WARNING_FILE.parent.mkdir(parents=True, exist_ok=True)
        with RESOURCES_NOT_FOUND_WARNING_FILE.open("a", encoding="utf-8") as fp:
            fp.write(f"{line}\n")
    except Exception as exc:
        log_warn(
            f"写入 resources_not_found_warning.txt 失败 line={line!r} err={exc}"
        )


def _looks_like_not_found_body(text: str, *, content_type: str = "") -> bool:
    body = (text or "").strip().lower()
    if not body:
        return False

    # CSS / JS assets often contain numeric literals like "404"; do not treat that as 404 pages.
    ct = (content_type or "").lower()
    looks_like_html = (
        "<html" in body
        or "<!doctype" in body
        or "<body" in body
        or "<title" in body
    )
    if ("text/css" in ct or "javascript" in ct) and not looks_like_html:
        return False

    compact = re.sub(r"\s+", " ", body)
    if re.search(r"<title>\s*(404|not found|页面不存在)", compact):
        return True
    if re.search(r"\b404\b.{0,48}\b(not\s*found|error|页面|不存在|找不到)\b", compact):
        return True
    for k in (
        "404 not found",
        "404 page not found",
        "resource not found",
        "页面不存在",
        "您访问的页面不存在",
        "找不到该页面",
    ):
        if k in compact:
            return True
    return False


def _is_textual_content_type(content_type: str) -> bool:
    ct = (content_type or "").lower()
    if not ct:
        return False
    return (
        ct.startswith("text/")
        or "json" in ct
        or "xml" in ct
        or "javascript" in ct
        or "x-www-form-urlencoded" in ct
    )


def _looks_like_expired_signature_body(text: str) -> bool:
    body = (text or "").lower()
    if not body:
        return False
    return (
        ("request has expired" in body)
        or ("<code>accessdenied</code>" in body and "expires" in body)
    )


def _text_decode_quality_score(text: str) -> int:
    if not text:
        return -10_000
    lower = text.lower()
    cjk = len(re.findall(r"[\u4e00-\u9fff]", text))
    bad_ctrl = len(re.findall(r"[\x00-\x08\x0B\x0C\x0E-\x1F\x7F-\x9F]", text))
    replacement = text.count("\ufffd")
    moji = sum(
        text.count(t)
        for t in ("锟", "鍛", "鐩", "鎴", "鏄", "銆", "鈥", "锛", "氓", "聽", "芒聙")
    )
    angle = text.count("<")
    if angle < 8 and ("<html" not in lower and "<body" not in lower):
        return -50_000 - bad_ctrl * 40 - replacement * 30
    html_bonus = 0
    if "<html" in lower:
        html_bonus += 220
    if "<head" in lower:
        html_bonus += 120
    if "<body" in lower:
        html_bonus += 120
    if "</" in lower:
        html_bonus += 80
    html_bonus += min(400, angle * 2)
    if angle < 8:
        html_bonus -= 1000
    return cjk * 2 - bad_ctrl * 40 - replacement * 30 - moji * 4 + html_bonus


def _decode_html_bytes(raw: bytes, preferred_encoding: str | None = None) -> str:
    if not raw:
        return ""
    def _norm_enc(enc: str) -> str:
        e = (enc or "").strip().lower().replace("_", "-")
        if e in ("utf8",):
            return "utf-8"
        if e in ("gb2312", "gbk"):
            return "gb18030"
        return e

    candidates: list[str] = []
    preferred_norm = _norm_enc(preferred_encoding or "")
    if preferred_norm:
        candidates.append(preferred_norm)
    meta_norm = ""
    m = HTML_CHARSET_META_RE.search(raw[:4096])
    if m:
        try:
            meta_norm = _norm_enc(m.group(1).decode("ascii", errors="ignore"))
            if meta_norm:
                candidates.append(meta_norm)
        except Exception:
            pass

    # Strongly trust explicit charset (HTTP / meta) when it decodes strictly and looks like HTML.
    for enc in [preferred_norm, meta_norm]:
        if not enc:
            continue
        try:
            txt = raw.decode(enc, errors="strict")
        except Exception:
            continue
        lower = txt.lower()
        if "<html" in lower or "<!doctype" in lower or "<body" in lower:
            return txt

    candidates.extend(["utf-8", "gb18030", "gbk", "utf-16"])

    best = ""
    best_score = -10**9
    seen: set[str] = set()
    for enc in candidates:
        e = (enc or "").strip().lower()
        if not e or e in seen:
            continue
        seen.add(e)
        try:
            txt = raw.decode(e)
        except Exception:
            txt = raw.decode(e, errors="replace")
        sc = _text_decode_quality_score(txt)
        if sc > best_score:
            best_score = sc
            best = txt
    if best:
        return best
    return raw.decode("utf-8", errors="replace")


def _response_text_html(resp: requests.Response) -> str:
    enc = (resp.encoding or "").strip().lower()
    preferred = None if enc in ("", "iso-8859-1", "latin-1") else resp.encoding
    return _decode_html_bytes(resp.content or b"", preferred_encoding=preferred)


def _filename_from_content_disposition(content_disposition: str) -> str:
    cd = (content_disposition or "").strip()
    if not cd:
        return ""
    m_star = re.search(r"filename\*\s*=\s*([^;]+)", cd, re.I)
    if m_star:
        raw = m_star.group(1).strip().strip('"')
        # RFC 5987: UTF-8''encoded-name
        if "''" in raw:
            raw = raw.split("''", 1)[1]
        raw = unquote(raw)
        return raw.strip()
    m = re.search(r'filename\s*=\s*"([^"]+)"', cd, re.I)
    if m:
        return m.group(1).strip()
    m2 = re.search(r"filename\s*=\s*([^;]+)", cd, re.I)
    if m2:
        return m2.group(1).strip().strip('"')
    return ""


def _suffix_from_document_magic(data: bytes) -> str | None:
    if not data:
        return None
    if data.startswith(b"%PDF-"):
        return ".pdf"
    if data.startswith(b"PK\x03\x04"):
        # OOXML family: docx/pptx/xlsx; preview downloads in this pipeline are commonly docx.
        return ".docx"
    if data.startswith(b"\xD0\xCF\x11\xE0\xA1\xB1\x1A\xE1"):
        # OLE family: doc/ppt/xls; default to doc for word-only alignment.
        return ".doc"
    return None


def fetch_download_document_url(
    session: requests.Session, article_id: str, *, source_url: str
) -> str | None:
    aid = (article_id or "").strip()
    if not aid or not aid.isdigit():
        return None
    ua = session.headers.get("User-Agent", "Mozilla/5.0")
    headers = {"Referer": source_url, "User-Agent": ua}
    info_url = f"http://www.360doc.com/ajax/getdownloadinfo.ashx?articleid={aid}"
    info_resp = request_with_retry(session, info_url, headers=headers, timeout=TIMEOUT)
    info = (info_resp.text or "").strip()
    if info not in {"1", "2"}:
        return None
    doc_url = f"http://www.360doc.com/ajax/getdownloaddocument.ashx?articleid={aid}"
    doc_resp = request_with_retry(session, doc_url, headers=headers, timeout=TIMEOUT)
    raw = (doc_resp.text or "").strip()
    if not raw or raw == "-1":
        return None
    return raw if raw.startswith(("http://", "https://")) else f"http://{raw}"


def download_original_preview_document(
    session: requests.Session,
    *,
    article_id: str,
    source_url: str,
    output_base_path: Path,
) -> Path | None:
    download_url = fetch_download_document_url(
        session, article_id, source_url=source_url
    )
    if not download_url:
        return None
    ua = session.headers.get("User-Agent", "Mozilla/5.0")
    headers = {"Referer": source_url, "User-Agent": ua}
    resp = request_with_retry(session, download_url, headers=headers, timeout=TIMEOUT)
    data = resp.content or b""
    if not data:
        return None

    cd_name = _filename_from_content_disposition(
        resp.headers.get("Content-Disposition", "")
    )
    ext = Path(cd_name).suffix.lower() if cd_name else ""
    if not ext:
        ext = Path(urlparse(resp.url).path).suffix.lower()
    if not ext:
        ext = _suffix_from_document_magic(data) or ".bin"

    out = output_base_path.with_suffix(ext)
    out.parent.mkdir(parents=True, exist_ok=True)
    out.write_bytes(data)
    return out


def _is_preview_only_article_html(
    *, soup: BeautifulSoup, raw_html: str, source_url: str
) -> bool:
    if extract_body_tag_standard(soup) is not None:
        return False
    if parse_word_document_meta(raw_html):
        return True
    if parse_pdf_image_urls(raw_html, source_url):
        return True
    if parse_ppt_image_urls(raw_html, source_url):
        return True
    return False


def _is_preview_content_node(node: Tag | None) -> bool:
    if not isinstance(node, Tag):
        return False
    cls = set(node.get("class") or [])
    return bool(
        {"word-document-preview", "pdf-document-preview", "ppt-document-preview"} & cls
    )


def _is_preview_clean_html_file(path: Path) -> bool:
    if not path.is_file():
        return False
    text = path.read_text(encoding="utf-8", errors="ignore")
    return any(
        k in text
        for k in (
            "word-document-preview",
            "pdf-document-preview",
            "ppt-document-preview",
        )
    )


def try_direct_download_preview_document_for_word_only(
    path: Path,
    session: requests.Session,
    *,
    force: bool,
) -> tuple[str, Path | None]:
    article_id = extract_article_id(path)
    source_url = guess_source_url(path)
    raw = _decode_html_bytes(path.read_bytes())
    soup = BeautifulSoup(raw, "html.parser")
    if not _is_preview_only_article_html(soup=soup, raw_html=raw, source_url=source_url):
        return "not_preview", None
    out_base = path.with_suffix("")
    # Respect incremental mode: if any aligned output already exists, skip.
    if not force:
        for ext in (".docx", ".doc", ".pdf", ".ppt", ".pptx"):
            cand = out_base.with_suffix(ext)
            if cand.is_file():
                return "skipped", cand
    try:
        saved = download_original_preview_document(
            session,
            article_id=article_id,
            source_url=source_url,
            output_base_path=out_base,
        )
    except Exception as exc:
        log_warn(f"原文档直下载失败 {path.name}: {exc}")
        return "failed", None
    if saved is None:
        return "failed", None
    return "processed", saved


def _category_dir_id_from_name(article_dir_name: str) -> str:
    m = re.match(r"^(-?\d+)-", (article_dir_name or "").strip())
    if not m:
        return ""
    raw = m.group(1)
    if raw.startswith("-"):
        raw = raw[1:]
    return raw


def _domain_hint_from_article_dir(article_dir_name: str) -> str:
    did = _category_dir_id_from_name(article_dir_name)
    if not did:
        return ""
    return _CATEGORY_ARTNUM_BY_DIR_ID.get(did, "")


def sanitize_name(name: str, fallback: str) -> str:
    name = INVALID_NAME_RE.sub("_", name).strip().rstrip(".")
    return name or fallback


def _is_html_inside_clean_resource_subdir(path: Path) -> bool:
    # HTML files under clean_<stem>/ are resource fragments and are excluded from article scans.
    for parent in path.parents:
        pn = parent.name
        if pn.startswith(CLEAN_HTML_PREFIX) and not pn.lower().endswith(".html"):
            return True
    return False


def iter_library_article_html_files(root: Path) -> list[Path]:
    # Candidate article files: id-prefixed raw HTML, or orphan clean_ files without matching raw files.
    # Exclude HTML fragments inside clean_* resource subdirectories (for example res_2.html).
    out: list[Path] = []
    pfx = CLEAN_HTML_PREFIX
    lpfx = pfx.lower()
    for p in root.rglob("*.html"):
        if not p.is_file() or _is_html_inside_clean_resource_subdir(p):
            continue
        name = p.name
        ln = name.lower()
        if ln.startswith(lpfx):
            raw_name = name[len(pfx) :]
            if p.with_name(raw_name).is_file():
                continue
            out.append(p)
            continue
        if re.match(r"^\d+-", name):
            out.append(p)
    out.sort()
    return out


def normalize_url(raw: str, base_url: str) -> str:
    if not raw:
        return ""
    raw = raw.strip()
    if raw.startswith(("javascript:", "#", "mailto:", "tel:", "data:")):
        return ""
    if raw.startswith("//"):
        return "http:" + raw
    return urljoin(base_url, raw)


def is_localizable_url(url: str) -> bool:
    if not url:
        return False
    parsed = urlparse(url)
    if parsed.scheme not in ("http", "https"):
        return False
    host = parsed.netloc.lower()
    return any(host.endswith(domain) for domain in ALLOWED_HOSTS)


def request_with_retry(
    session: requests.Session,
    url: str,
    headers: dict | None = None,
    *,
    timeout: int = TIMEOUT,
    retries: int = MAX_RETRY,
    use_session_cookies: bool = True,
    bypass_env_proxy: bool = False,
) -> requests.Response:
    last_exc: Exception | None = None
    for attempt in range(1, retries + 2):
        try:
            req_kwargs: dict = {"timeout": timeout, "headers": headers}
            if bypass_env_proxy:
                # 某些环境变量代理会把 360doc 资源图链放大成 502，资源下载时可按需直连。
                req_kwargs["proxies"] = {"http": None, "https": None}
            if use_session_cookies:
                resp = session.get(url, **req_kwargs)
            else:
                # 资源直链尽量不带 Cookie，贴近浏览器跨域图片请求。
                resp = requests.get(url, **req_kwargs)
            if resp.status_code == 403 and _looks_like_expired_signature_body(resp.text):
                raise ResourceExpiredError(f"resource signature expired url={url}")
            resp_ct = resp.headers.get("content-type", "")
            if resp.status_code == 404 or (
                resp.status_code < 500
                and _is_textual_content_type(resp_ct)
                and _looks_like_not_found_body(resp.text, content_type=resp_ct)
            ):
                raise ResourceNotFoundError(f"resource not found url={url}")
            if resp.status_code in TRANSIENT_GATEWAY_STATUS_CODES:
                if _looks_like_expired_signature_body(resp.text):
                    raise ResourceExpiredError(f"resource signature expired url={url}")
                raise ResourceGatewayError(
                    f"{resp.status_code} Server Error: Bad Gateway for url: {url}"
                )
            resp.raise_for_status()
            return resp
        except Exception as exc:
            if isinstance(exc, ResourceGatewayError):
                last_exc = exc
                if attempt < (retries + 1):
                    time.sleep(random.uniform(*TRANSIENT_GATEWAY_RETRY_SLEEP_SEC))
                    continue
                break
            if isinstance(exc, requests.HTTPError) and exc.response is not None:
                status = exc.response.status_code
                if status in RATE_LIMIT_STATUS_CODES:
                    preview = (exc.response.text or "")[:300]
                    raise CleanRateLimitError(
                        f"clean resource request blocked url={url} status={status} body={preview!r}"
                    ) from exc
            last_exc = exc
    assert last_exc is not None
    raise last_exc

def extract_article_meta(soup: BeautifulSoup) -> tuple[str, str, str]:
    title = ""
    title_node = soup.select_one("#GLTitile")
    if title_node:
        title = title_node.get_text(strip=True)
    if not title:
        h1 = soup.select_one("h1#titiletext")
        if h1:
            title = h1.get_text(" ", strip=True)
    if not title and soup.title:
        title = soup.title.get_text(strip=True)

    author = ""
    author_node = soup.select_one("#savernickname")
    if author_node:
        author = author_node.get_text(strip=True)

    publish_date = ""
    top_data = soup.select_one(".art_topdata")
    if top_data:
        date_match = DATE_RE.search(top_data.get_text(" ", strip=True))
        if date_match:
            publish_date = date_match.group(0)

    return title or "无标题", author or "未知发布者", publish_date or ""


def extract_body_tag_standard(soup: BeautifulSoup) -> Tag | None:
    content_node = soup.select_one("#artContent")
    if content_node is not None:
        return content_node
    n = soup.select_one("#printArticle")
    if n is not None:
        return n
    return soup.select_one("#content")


def parse_word_document_meta(raw_html: str) -> tuple[str, int] | None:
    m_url = WORD_META_WORDURL_RE.search(raw_html)
    m_pages = WORD_META_PAGENUM_RE.search(raw_html)
    if not m_url or not m_pages:
        return None
    base = m_url.group(1).strip().rstrip("/")
    n = int(m_pages.group(1))
    if n < 1:
        return None
    if n > MAX_WORD_PREVIEW_PAGES:
        log_warn(
            f"Word 预览页数过大，截断为前 {MAX_WORD_PREVIEW_PAGES} 页（原始 {n} 页）"
        )
        n = MAX_WORD_PREVIEW_PAGES
    return base, n


def parse_word_css_url(raw_html: str, word_base: str) -> str:
    m_css = WORD_META_WORDCSSURL_RE.search(raw_html or "")
    if m_css:
        css_base = (m_css.group(1) or "").strip().rstrip("/")
        if css_base:
            return f"{css_base}.css"
    return f"{word_base}.css"


def parse_word_preview_css_urls(
    *, soup: BeautifulSoup, raw_html: str, source_url: str, word_base: str
) -> list[str]:
    out: list[str] = []
    seen: set[str] = set()

    def _add(u: str) -> None:
        uu = normalize_url(u, source_url)
        uu = _prefer_working_360doc_image_host(uu)
        if not is_localizable_url(uu):
            return
        if uu in seen:
            return
        seen.add(uu)
        out.append(uu)

    _add(parse_word_css_url(raw_html, word_base))
    for link in soup.select("link[rel][href]"):
        if not isinstance(link, Tag):
            continue
        rel = " ".join(link.get("rel", [])).lower() if link.get("rel") else ""
        if "stylesheet" not in rel:
            continue
        href = str(link.get("href", "")).strip()
        if not href:
            continue
        hlow = href.lower()
        if (
            "docartpage" in hlow
            or "page_word" in hlow
            or "wordbase" in hlow
            or "wordfancy" in hlow
        ):
            _add(href)
    return out


def fetch_word_preview_body(
    session: requests.Session,
    word_base: str,
    page_count: int,
    source_url: str,
    article_id: str,
    article_title: str,
    article_dir_name: str,
) -> Tag | None:
    wrapper_soup = BeautifulSoup(
        '<div class="word-document-preview"></div>', "html.parser"
    )
    wrapper = wrapper_soup.div
    if wrapper is None:
        return None

    ua = session.headers.get("User-Agent", "Mozilla/5.0")
    headers = {"Referer": source_url, "User-Agent": ua}
    any_page_ok = False
    ok_pages = 0
    fail_pages = 0
    consecutive_failures = 0
    started_at = time.time()
    last_heartbeat_at = started_at

    for p in range(1, page_count + 1):
        page_url = f"{word_base}_{p}.html"
        page_div = wrapper_soup.new_tag(
            "div", attrs={"class": "word-preview-page", "data-page": str(p)}
        )
        try:
            resp = request_with_retry(
                session,
                page_url,
                headers=headers,
                timeout=int(max(1.0, WORD_PREVIEW_REQUEST_TIMEOUT)),
                retries=max(0, WORD_PREVIEW_REQUEST_RETRIES),
            )
            sub = BeautifulSoup(_response_text_html(resp), "html.parser")
            for bad in sub.select("script,style,noscript,iframe"):
                bad.decompose()
            src_body = sub.body
            if src_body:
                for ch in list(src_body.children):
                    page_div.append(ch.extract())
            else:
                root = sub.find(["div", "img", "section", "article", "p", "table"])
                if root is not None:
                    page_div.append(root.extract())
            if page_div.get_text(strip=True) or page_div.find(["img", "a", "table"]):
                any_page_ok = True
                ok_pages += 1
                consecutive_failures = 0
            wrapper.append(page_div)
        except CleanRateLimitError:
            raise
        except ResourceNotFoundError:
            fail_pages += 1
            consecutive_failures += 1
            append_resource_not_found_warning_line(
                f"article_id={article_id}\tarticle={article_title}\tdir={article_dir_name}"
                f"\tresource={page_url}\tnot_found=1"
            )
            wrapper.append(page_div)
        except Exception as exc:
            fail_pages += 1
            consecutive_failures += 1
            log_warn(f"Word 预览页拉取失败 art={article_id} url={page_url} err={exc}")
            append_clean_resource_failure_line(
                article_id=article_id,
                article_title=article_title,
                article_dir_name=article_dir_name,
                resource_url=page_url,
                error=exc,
            )
            wrapper.append(page_div)

        now = time.time()
        should_log = (
            p == page_count
            or p % WORD_PREVIEW_PROGRESS_EVERY_PAGES == 0
            or (now - last_heartbeat_at) >= WORD_PREVIEW_HEARTBEAT_SEC
        )
        if should_log:
            elapsed = max(0.001, now - started_at)
            rate = p / elapsed
            log_info(
                f"Word 预览进度 art={article_id} {p}/{page_count} "
                f"ok={ok_pages} fail={fail_pages} rate={rate:.2f}p/s"
            )
            last_heartbeat_at = now

        if (
            WORD_PREVIEW_MAX_CONSECUTIVE_FAILURES > 0
            and consecutive_failures >= WORD_PREVIEW_MAX_CONSECUTIVE_FAILURES
        ):
            log_warn(
                f"Word 预览连续失败过多，提前终止 art={article_id} "
                f"consecutive_fail={consecutive_failures} page={p}/{page_count}"
            )
            break
        time.sleep(random.uniform(*WORD_PREVIEW_PAGE_SLEEP_SEC))

    if not any_page_ok:
        return None
    return wrapper


def parse_ppt_image_urls(raw_html: str, source_url: str) -> list[str]:
    m = PPT_IMG_ARR_RE.search(raw_html or "")
    if not m:
        return []
    block = m.group(1) or ""
    out: list[str] = []
    seen: set[str] = set()
    for um in PPT_IMG_URL_RE.finditer(block):
        raw = (um.group(1) or "").strip()
        if not raw:
            continue
        abs_url = normalize_url(raw, source_url)
        if not is_localizable_url(abs_url):
            continue
        abs_url = _prefer_working_360doc_image_host(abs_url)
        if abs_url in seen:
            continue
        seen.add(abs_url)
        out.append(abs_url)
    return out


def _collect_urls_from_js_array(block: str, source_url: str) -> list[str]:
    out: list[str] = []
    seen: set[str] = set()
    for um in PPT_IMG_URL_RE.finditer(block or ""):
        raw = (um.group(1) or "").strip()
        if not raw:
            continue
        abs_url = normalize_url(raw, source_url)
        if not is_localizable_url(abs_url):
            continue
        abs_url = _prefer_working_360doc_image_host(abs_url)
        if abs_url in seen:
            continue
        seen.add(abs_url)
        out.append(abs_url)
    return out


def _parse_pdf_urls_from_ger_looking_user_info(
    raw_html: str, source_url: str
) -> list[str]:
    gm = GER_LOOKING_USER_INFO_RE.search(raw_html or "")
    if not gm:
        return []
    args = gm.group(1) or ""
    page_match = re.search(r",\s*(\d+)\s*,\s*\d+\s*,\s*\d+\s*$", args)
    first_match = PDF_FIRST_PAGE_URL_RE.search(args)
    if not page_match or not first_match:
        return []
    try:
        page_count = int(page_match.group(1))
    except Exception:
        return []
    if page_count < 1:
        return []
    first_url = (first_match.group(1) or "").strip()
    if not first_url:
        return []
    index = first_match.group(2)
    first_abs = normalize_url(first_url, source_url)
    m = re.search(rf"^(.*_){re.escape(index)}(\.[^.?#]+)(.*)$", first_abs, re.I)
    if not m:
        return []
    prefix, suffix, tail = m.group(1), m.group(2), m.group(3)
    out: list[str] = []
    for p in range(1, page_count + 1):
        u = f"{prefix}{p}{suffix}{tail}"
        if not is_localizable_url(u):
            continue
        out.append(_prefer_working_360doc_image_host(u))
    return out


def parse_pdf_image_urls(raw_html: str, source_url: str) -> list[str]:
    m = PDF_IMG_ARR_RE.search(raw_html or "")
    if m:
        urls = _collect_urls_from_js_array(m.group(1) or "", source_url)
        if urls:
            return urls
    return _parse_pdf_urls_from_ger_looking_user_info(raw_html, source_url)


def build_pdf_preview_body(pdf_urls: list[str]) -> Tag | None:
    if not pdf_urls:
        return None
    wrapper_soup = BeautifulSoup(
        '<div class="pdf-document-preview"></div>', "html.parser"
    )
    wrapper = wrapper_soup.div
    if wrapper is None:
        return None
    for idx, u in enumerate(pdf_urls, start=1):
        page_div = wrapper_soup.new_tag(
            "div", attrs={"class": "pdf-preview-page", "data-page": str(idx)}
        )
        img = wrapper_soup.new_tag(
            "img",
            attrs={"src": u, "alt": f"pdf-page-{idx}", "data-pdf-page": str(idx)},
        )
        page_div.append(img)
        wrapper.append(page_div)
    return wrapper


def build_ppt_preview_body(ppt_urls: list[str]) -> Tag | None:
    if not ppt_urls:
        return None
    wrapper_soup = BeautifulSoup(
        '<div class="ppt-document-preview"></div>', "html.parser"
    )
    wrapper = wrapper_soup.div
    if wrapper is None:
        return None
    for idx, u in enumerate(ppt_urls, start=1):
        page_div = wrapper_soup.new_tag(
            "div", attrs={"class": "ppt-preview-page", "data-page": str(idx)}
        )
        img = wrapper_soup.new_tag(
            "img",
            attrs={"src": u, "alt": f"ppt-page-{idx}", "data-ppt-page": str(idx)},
        )
        page_div.append(img)
        wrapper.append(page_div)
    return wrapper


def resolve_content_node(
    *,
    soup: BeautifulSoup,
    raw_html: str,
    session: requests.Session,
    source_url: str,
    article_id: str,
    article_title: str,
    article_dir_name: str,
) -> Tag | None:
    content = extract_body_tag_standard(soup)
    if content is not None:
        return content
    pdf_urls = parse_pdf_image_urls(raw_html, source_url)
    pdf_content = build_pdf_preview_body(pdf_urls)
    if pdf_content is not None:
        log_info(f"PDF 预览正文 art={article_id} pages={len(pdf_urls)}")
        return pdf_content
    word_meta = parse_word_document_meta(raw_html)
    if word_meta:
        word_base, page_count = word_meta
        log_info(
            f"Word 预览正文 art={article_id} pages={page_count} base={word_base}"
        )
        word_css_urls = parse_word_preview_css_urls(
            soup=soup, raw_html=raw_html, source_url=source_url, word_base=word_base
        )
        content = fetch_word_preview_body(
            session,
            word_base,
            page_count,
            source_url,
            article_id,
            article_title,
            article_dir_name,
        )
        if content is not None:
            content["data-word-css-urls"] = "||".join(word_css_urls)
            return content
    ppt_urls = parse_ppt_image_urls(raw_html, source_url)
    ppt_content = build_ppt_preview_body(ppt_urls)
    if ppt_content is not None:
        log_info(f"PPT 预览正文 art={article_id} pages={len(ppt_urls)}")
        return ppt_content
    return None


def build_clean_soup(title: str, author: str, publish_date: str, content_node: Tag) -> BeautifulSoup:
    # Build cleaned DOM: article card with title, metadata, and #content body; omit #artContent shell.
    extra_head_links: list[str] = []
    is_word_preview = "word-document-preview" in (content_node.get("class") or [])
    if is_word_preview:
        raw_urls = str(content_node.get("data-word-css-urls", "")).strip()
        if raw_urls:
            for u in raw_urls.split("||"):
                uu = u.strip()
                if uu and uu not in extra_head_links:
                    extra_head_links.append(uu)

    links_html = "\n".join(
        f'  <link rel="stylesheet" href="{u}" data-doc360-localize="1"/>' for u in extra_head_links
    )
    if links_html:
        links_html = links_html + "\n"
    if is_word_preview:
        clean_html = f"""<!doctype html>
<html lang="zh-CN">
<head>
  <meta charset="utf-8">
  <meta name="viewport" content="width=device-width, initial-scale=1">
  <title>{title}</title>
  <style>
    html,body{{margin:0;padding:0;background:#f3f4f6;overflow-x:hidden;}}
    body{{box-sizing:border-box;padding:12px 14px;text-align:center;}}
    .doc360-card{{display:inline-block;vertical-align:top;width:fit-content;max-width:calc(100vw - 28px);margin:0 auto;background:#fff;border-radius:0.5px;box-shadow:0 2px 14px rgba(0,0,0,0.07);padding:18px 22px 22px;box-sizing:border-box;overflow:visible;text-align:left;}}
    #content{{display:block;width:fit-content;max-width:none;margin:0 auto;padding:0;box-sizing:border-box;text-align:left;}}
    #content img{{max-width:none;height:auto;display:inline;margin:0;}}
    .word-document-preview{{display:block;width:fit-content;max-width:none;overflow:visible !important;margin:0 auto !important;padding:0 !important;}}
    .word-document-preview,.word-preview-page,.pf{{margin-left:auto !important;margin-right:auto !important;}}
    .word-document-preview .word-preview-page{{display:block;width:fit-content !important;max-width:none !important;margin:0 auto !important;padding:0 !important;border:0 !important;box-shadow:none !important;float:none !important;overflow:visible !important;}}
    .word-document-preview .word-preview-page:first-child{{margin-top:0 !important;}}
    .word-document-preview .pf{{float:none !important;margin:0 auto !important;border:0 !important;box-shadow:none !important;}}
    .word-document-preview .pc,.word-document-preview .c{{border:0 !important;box-shadow:none !important;}}
  </style>
{links_html}</head>
<body>
  <article class="doc360-card"><div id="content"></div></article>
</body>
</html>"""
    else:
        clean_html = f"""<!doctype html>
<html lang="zh-CN">
<head>
  <meta charset="utf-8">
  <meta name="viewport" content="width=device-width, initial-scale=1">
  <title>{title}</title>
  <style>
    body{{margin:0;padding:12px 14px;font-family:Arial,"Microsoft YaHei",sans-serif;line-height:1.75;background:#f3f4f6;}}
    article.doc360-card{{max-width:980px;margin:0 auto;background:#fff;border-radius:0.5px;box-shadow:0 2px 14px rgba(0,0,0,0.07);padding:22px 24px 28px;box-sizing:border-box;display:flex;flex-direction:column;align-items:stretch;}}
    .doc360-title{{align-self:center;width:fit-content;max-width:100%;text-align:center;}}
    .doc360-title h1{{margin:0;font-size:30px;line-height:1.35;font-weight:normal;}}
    .doc360-meta{{align-self:center;width:fit-content;max-width:100%;margin:10px 0 18px;color:#555;font-size:14px;text-align:center;}}
    #content.doc360-main{{width:100%;max-width:100%;box-sizing:border-box;}}
    #content.doc360-main img{{max-width:100%;height:auto;display:block;margin-left:auto;margin-right:auto;}}
    #content.doc360-main table{{max-width:100%;margin:0;}}
    #content.doc360-main .word-document-preview img{{max-width:none !important;display:inline;margin:0;}}
    #content.doc360-main .word-document-preview{{overflow:auto;}}
  </style>
{links_html}</head>
<body>
  <article class="doc360-card">
    <div class="doc360-title"><h1 id="title"></h1></div>
    <div class="doc360-meta"><span id="author"></span> <span id="date"></span></div>
    <div class="doc360-main" id="content"></div>
  </article>
</body>
</html>"""
    out = BeautifulSoup(clean_html, "html.parser")
    if not is_word_preview:
        t_el = out.select_one("#title")
        if t_el:
            t_el.string = title
        a_el = out.select_one("#author")
        if a_el:
            a_el.string = author
        d_el = out.select_one("#date")
        if d_el:
            d_el.string = publish_date
    content_target = out.select_one("#content")
    if content_target is None:
        raise ValueError("清洗模板缺少 #content")
    content_copy = BeautifulSoup(str(content_node), "html.parser")
    root_content = content_copy.find(True) or content_copy
    if isinstance(root_content, Tag):
        root_content.attrs.pop("data-word-css-url", None)
        root_content.attrs.pop("data-word-css-urls", None)
    for bad in root_content.select("script,style,iframe,noscript"):
        bad.decompose()
    if isinstance(root_content, Tag) and (root_content.name or "").lower() in ("td", "th"):
        root_content.name = "div"
    if isinstance(root_content, Tag):
        rid = (root_content.get("id") or "").strip()
        if rid in _CONTENT_WRAPPER_IDS:
            for ch in list(root_content.children):
                content_target.append(ch)
        else:
            content_target.append(root_content)
    return out


_IMG_PLACEHOLDER_SRC_RE = re.compile(
    r"(?:space|blank|spacer|transparent|1x1|default|loading)\.(?:gif|png)|pixel\.gif$",
    re.I,
)


def _prefer_working_360doc_image_host(url: str) -> str:
    # data360-src often points to checki* hosts; rewrite to image* variant to avoid gohost 404 hops.
    if not url:
        return url
    try:
        parts = urlparse(url)
        host = (parts.netloc or "").lower()
        if "checki" in host and host.endswith("360doc.com"):
            new_netloc = host.replace("checki", "image", 1)
            return urlunparse(
                (
                    parts.scheme or "http",
                    new_netloc,
                    parts.path,
                    parts.params,
                    parts.query,
                    parts.fragment,
                )
            )
    except Exception:
        pass
    return url


def _strip_url_query(url: str) -> str:
    if not url:
        return url
    try:
        p = urlparse(url)
        return urlunparse((p.scheme, p.netloc, p.path, p.params, "", ""))
    except Exception:
        return url


def _https_variant(url: str) -> str:
    uu = (url or "").strip()
    if not uu.startswith("http://"):
        return ""
    try:
        h = (urlparse(uu).netloc or "").lower()
        # 这类 360doc 老 CDN 节点经常 HTTPS EOF，避免优先走 https。
        if h.endswith(".360doc.com") and (
            h.startswith("imgu")
            or h.startswith("imgi")
            or h.startswith("checku")
            or h.startswith("checki")
        ):
            return ""
    except Exception:
        pass
    return "https://" + uu[7:]


def _prefer_legacy_360doc_http(url: str) -> str:
    uu = (url or "").strip()
    if not uu:
        return uu
    try:
        p = urlparse(uu)
        h = (p.netloc or "").lower()
        if p.scheme == "https" and h.endswith(".360doc.com") and (
            h.startswith("imgu")
            or h.startswith("imgi")
            or h.startswith("checku")
            or h.startswith("checki")
        ):
            return urlunparse(("http", p.netloc, p.path, p.params, p.query, p.fragment))
    except Exception:
        return uu
    return uu
    return ""


def _url_path_key(url: str) -> str:
    try:
        return urlparse(url).path or ""
    except Exception:
        return ""


def _rewrite_url_host(url: str, new_host: str) -> str:
    try:
        p = urlparse(url)
        return urlunparse((p.scheme, new_host, p.path, p.params, p.query, p.fragment))
    except Exception:
        return url


def _legacy_360doc_host_family(host: str) -> list[str]:
    h = (host or "").lower().strip()
    if not h.endswith(".360doc.com"):
        return []
    m = re.match(r"^(check|img)([ui])([0-9a-z]+)\.360doc\.com$", h)
    if not m:
        return [h]
    _, _, tail = m.groups()
    out = [
        f"checku{tail}.360doc.com",
        f"checki{tail}.360doc.com",
        f"imgu{tail}.360doc.com",
        f"imgi{tail}.360doc.com",
    ]
    uniq: list[str] = []
    for x in out:
        if x not in uniq:
            uniq.append(x)
    return uniq


def _build_article_signed_src_candidates(
    session: requests.Session, source_url: str
) -> tuple[dict[str, str], dict[str, str]]:
    # 从源文章页提取 data360-src -> src 的动态映射（src 常带 Expires/Signature/domain）。
    full_map: dict[str, str] = {}
    path_map: dict[str, str] = {}
    try:
        resp = session.get(
            source_url,
            timeout=TIMEOUT,
            headers={
                "Referer": BASE_URL + "/",
                "User-Agent": session.headers.get("User-Agent", "Mozilla/5.0"),
            },
        )
        if resp.status_code != 200:
            return full_map, path_map
        soup = BeautifulSoup(resp.text or "", "html.parser")
        for img in soup.find_all("img"):
            if not isinstance(img, Tag):
                continue
            raw_src = str(img.get("src", "")).strip()
            if not raw_src or raw_src.startswith("data:"):
                continue
            src_abs = normalize_url(raw_src, source_url)
            if not is_localizable_url(src_abs):
                continue
            src_abs = _prefer_working_360doc_image_host(src_abs)
            for key_attr in ("doc360img-src", "data360-src", "data-src", "data-original"):
                raw_key = str(img.get(key_attr, "")).strip()
                if not raw_key:
                    continue
                key_abs = _prefer_working_360doc_image_host(
                    normalize_url(raw_key, source_url)
                )
                if not key_abs:
                    continue
                full_map[key_abs] = src_abs
                kpath = _url_path_key(key_abs)
                if kpath:
                    path_map[kpath] = src_abs
    except Exception:
        return full_map, path_map
    return full_map, path_map


def _img_change_sign(params: dict[str, str]) -> str:
    parts: list[str] = []
    for k, v in params.items():
        sv = str(v)
        if sv == "":
            continue
        parts.append(f"{k}={sv}")
    parts.sort()
    s = "".join(parts)
    return hashlib.sha1(s.encode("utf-8")).hexdigest().upper()


def _request_changeurl_signed_images(
    session: requests.Session, source_url: str, raw_img_urls: list[str]
) -> tuple[dict[str, str], dict[str, str]]:
    urls = [(u or "").strip() for u in raw_img_urls if (u or "").strip()]
    if not urls:
        return {}, {}
    imgurl = ",".join(urls)
    payload = {"op": "changeurl", "imgurl": imgurl}
    sign = _img_change_sign(payload)
    try:
        resp = session.post(
            f"{IMG_CHANGE_URL}&_={int(time.time() * 1000)}",
            data={"imgurl": imgurl, "sign": sign},
            timeout=TIMEOUT,
            headers={
                "Referer": source_url,
                "User-Agent": session.headers.get("User-Agent", "Mozilla/5.0"),
                "X-Requested-With": "XMLHttpRequest",
            },
        )
    except Exception:
        return {}, {}
    if resp.status_code != 200:
        return {}, {}
    txt = (resp.text or "").strip()
    if not txt:
        return {}, {}

    try:
        data = json.loads(txt)
    except Exception:
        return {}, {}
    if str(data.get("status", "")).strip() != "1":
        return {}, {}
    raw_encoded = str(data.get("imgurl", "") or "").strip()
    if not raw_encoded:
        return {}, {}

    # 样本里是 URL 编码后的逗号串，单次解码后可按逗号切分。
    decoded = unquote(raw_encoded).replace("\\/", "/")
    out_list = [(u or "").strip() for u in decoded.split(",") if (u or "").strip()]

    by_input: dict[str, str] = {}
    by_path: dict[str, str] = {}
    now_ts = int(time.time())
    for idx, su in enumerate(out_list):
        abs_su = _prefer_working_360doc_image_host(normalize_url(su, source_url))
        if not is_localizable_url(abs_su):
            continue
        try:
            q = dict(parse_qsl(urlparse(abs_su).query, keep_blank_values=True))
            exp_raw = str(q.get("Expires", "")).strip()
            if exp_raw.isdigit() and int(exp_raw) <= now_ts + 15:
                # 已过期或即将过期的签名直接丢弃，避免“拿到即失效”。
                continue
        except Exception:
            pass
        if idx < len(urls):
            by_input[urls[idx]] = abs_su
        p = _url_path_key(abs_su)
        if p:
            by_path[p] = abs_su
    return by_input, by_path


def _img_delegated_to_parent_download_anchor(img: Tag) -> bool:
    # If parent <a href> is already a 360doc image link, skip the nested <img> to avoid duplicate fetches.
    for attr in ("doc360img-src", "data360-src", "data-src", "data-original"):
        raw_attr = str(img.get(attr, "")).strip()
        if raw_attr and not raw_attr.startswith("data:"):
            return False
    p = img.parent
    if not isinstance(p, Tag) or (p.name or "").lower() != "a":
        return False
    href = str(p.get("href", "")).strip()
    if not href or href.startswith(("javascript:", "#")):
        return False
    hlow = href.lower()
    return "360doc.com" in hlow and "downloadimg" in hlow


def _anchor_wraps_360doc_download_img(anchor: Tag) -> bool:
    if (anchor.name or "").lower() != "a":
        return False
    href = str(anchor.get("href", "")).strip()
    if not href:
        return False
    hlow = href.lower()
    if "360doc.com" not in hlow or "downloadimg" not in hlow:
        return False
    for img in anchor.find_all("img", recursive=True):
        if isinstance(img, Tag):
            return True
    return False


def _img_download_attr_name(tag: Tag) -> str | None:
    # Download attribute priority: explicit 360doc image attributes first, then src.
    for attr in ("doc360img-src", "data360-src", "data-src", "data-original", "src"):
        raw = str(tag.get(attr, "")).strip()
        if not raw or raw.startswith("data:"):
            continue
        if attr == "src" and _IMG_PLACEHOLDER_SRC_RE.search(raw):
            continue
        return attr
    return None


def collect_resource_nodes(soup: BeautifulSoup) -> list[tuple[Tag, str, str]]:
    # Return tuples: (tag, remote-url-attr, localized-path-attr). img always writes localized path to src.
    nodes: list[tuple[Tag, str, str]] = []
    head = soup.head
    if isinstance(head, Tag):
        for link in head.find_all("link", recursive=True):
            if not isinstance(link, Tag):
                continue
            if not link.has_attr("href"):
                continue
            rel = " ".join(link.get("rel", [])).lower() if link.get("rel") else ""
            if "stylesheet" not in rel:
                continue
            if str(link.get("data-doc360-localize", "")).strip() == "1":
                nodes.append((link, "href", "href"))
    root = soup.select_one("#content")
    if root is None:
        return nodes
    for tag in root.find_all(["a", "img", "source"], recursive=True):
        if not isinstance(tag, Tag):
            continue
        name = (tag.name or "").lower()
        if name == "a" and tag.has_attr("href"):
            if _anchor_wraps_360doc_download_img(tag):
                continue
            nodes.append((tag, "href", "href"))
        elif name == "source" and tag.has_attr("src"):
            nodes.append((tag, "src", "src"))
        elif name == "img":
            if _img_delegated_to_parent_download_anchor(tag):
                continue
            ra = _img_download_attr_name(tag)
            if ra:
                nodes.append((tag, ra, "src"))
    return nodes


def _suffix_from_content_type(content_type: str, fallback: str) -> str:
    ct = (content_type or "").lower().split(";", 1)[0].strip()
    if ct in ("image/jpeg", "image/jpg"):
        return ".jpeg"
    if ct == "image/png":
        return ".png"
    if ct == "image/gif":
        return ".gif"
    if ct == "image/webp":
        return ".webp"
    if ct == "image/bmp":
        return ".bmp"
    if ct == "image/svg+xml":
        return ".svg"
    return fallback


def _suffix_from_magic(data: bytes) -> str | None:
    if not data:
        return None
    if data.startswith(b"\xFF\xD8\xFF"):
        return ".jpeg"
    if data.startswith(b"\x89PNG\r\n\x1a\n"):
        return ".png"
    if data.startswith((b"GIF87a", b"GIF89a")):
        return ".gif"
    if data.startswith(b"BM"):
        return ".bmp"
    if len(data) >= 12 and data[:4] == b"RIFF" and data[8:12] == b"WEBP":
        return ".webp"
    if data.startswith(b"<svg") or data.startswith(b"<?xml"):
        # SVG 文件常以 XML 声明或 svg 标签开头。
        head = data[:200].lower()
        if b"<svg" in head:
            return ".svg"
    return None


def suffix_from_url(url: str, fallback: str) -> str:
    path = urlparse(url).path
    if "." in path:
        ext = path.rsplit(".", 1)[-1].lower()
        ext = re.sub(r"[^a-z0-9]", "", ext)
        if ext:
            return "." + ext[:8]
    return fallback


def localize_resources(
    clean_soup: BeautifulSoup,
    source_url: str,
    clean_output_path: Path,
    session: requests.Session,
    *,
    article_id: str,
    article_title: str,
    article_dir_name: str,
) -> ResourceLocalizationResult:
    def _rewrite_css_and_localize_deps(
        css_bytes: bytes, css_remote_url: str, css_local_name: str
    ) -> bytes:
        css_text = _decode_html_bytes(css_bytes)
        if not css_text.strip():
            return css_bytes
        css_dir = clean_output_path.with_suffix("")

        dep_cache: dict[str, str] = {}
        dep_seq = 0

        def _repl(m: re.Match[str]) -> str:
            nonlocal dep_seq
            quote = m.group(1) or ""
            raw = (m.group(2) or "").strip()
            if not raw or raw.startswith(("data:", "#", "javascript:")):
                return m.group(0)
            abs_dep = normalize_url(raw, css_remote_url)
            abs_dep = _prefer_working_360doc_image_host(abs_dep)
            if not is_localizable_url(abs_dep):
                return m.group(0)
            if abs_dep in dep_cache:
                local_ref = dep_cache[abs_dep]
                return f"url({quote}{local_ref}{quote})"
            try:
                dep_resp = request_with_retry(
                    session,
                    abs_dep,
                    headers={
                        "Referer": source_url,
                        "User-Agent": session.headers.get("User-Agent", "Mozilla/5.0"),
                    },
                    timeout=RESOURCE_REQUEST_TIMEOUT,
                    retries=RESOURCE_REQUEST_RETRIES,
                    use_session_cookies=("signature=" not in abs_dep.lower()),
                    bypass_env_proxy=True,
                )
                dep_data = dep_resp.content or b""
                if not dep_data:
                    return m.group(0)
                dep_ext = suffix_from_url(abs_dep, ".bin")
                dep_magic = _suffix_from_magic(dep_data)
                if dep_magic:
                    dep_ext = dep_magic
                elif dep_ext in (".bin", ".html", ".css"):
                    dep_ext = _suffix_from_content_type(
                        dep_resp.headers.get("content-type", ""),
                        dep_ext,
                    )
                dep_seq += 1
                dep_name = sanitize_name(
                    f"res_css_{Path(css_local_name).stem}_{dep_seq}{dep_ext}",
                    f"res_css_{dep_seq}{dep_ext}",
                )
                dep_path = css_dir / dep_name
                dep_path.write_bytes(dep_data)
                dep_cache[abs_dep] = dep_name
                return f"url({quote}{dep_name}{quote})"
            except Exception as exc:
                append_clean_resource_failure_line(
                    article_id=article_id,
                    article_title=article_title,
                    article_dir_name=article_dir_name,
                    resource_url=abs_dep,
                    error=exc,
                )
                return m.group(0)

        new_css = CSS_URL_FUNC_RE.sub(_repl, css_text)
        return new_css.encode("utf-8", errors="ignore")

    resource_nodes = collect_resource_nodes(clean_soup)
    if not resource_nodes:
        return ResourceLocalizationResult(downloaded=0, failed_urls=[])

    plan: list[tuple[Tag, str, str, str]] = []
    for tag, read_attr, write_attr in resource_nodes:
        raw = str(tag.get(read_attr, "")).strip()
        abs_url = normalize_url(raw, source_url)
        abs_url = _prefer_working_360doc_image_host(abs_url)
        if not is_localizable_url(abs_url):
            continue
        plan.append((tag, read_attr, write_attr, abs_url))

    if not plan:
        return ResourceLocalizationResult(downloaded=0, failed_urls=[])

    res_dir = clean_output_path.with_suffix("")
    downloaded = 0
    failed_urls: list[str] = []
    url_to_local: dict[str, str] = {}

    # 按首次出现顺序去重，保证 res_ 序号稳定且与历史行为一致。
    unique_urls: list[str] = []
    url_meta: dict[str, tuple[int, str]] = {}
    url_primary_tag: dict[str, Tag] = {}
    for tag, _, write_attr, abs_url in plan:
        if abs_url in url_meta:
            continue
        unique_urls.append(abs_url)
        file_seq = len(unique_urls)
        if (tag.name or "").lower() == "link":
            fallback_ext = ".css"
        else:
            fallback_ext = ".html" if write_attr == "href" else ".bin"
        url_meta[abs_url] = (file_seq, fallback_ext)
        url_primary_tag[abs_url] = tag

    signed_full_map: dict[str, str] | None = None
    signed_path_map: dict[str, str] | None = None
    changed_url_cache: dict[str, str] = {}
    changed_path_map: dict[str, str] = {}
    cache_lock = threading.Lock()
    domain_hint = _domain_hint_from_article_dir(article_dir_name)

    def _get_signed_maps() -> tuple[dict[str, str], dict[str, str]]:
        nonlocal signed_full_map, signed_path_map
        if signed_full_map is None or signed_path_map is None:
            signed_full_map, signed_path_map = _build_article_signed_src_candidates(
                session, source_url
            )
        return signed_full_map, signed_path_map

    def _candidate_urls_for_primary(url: str) -> list[str]:
        cands: list[str] = []
        now_ts = int(time.time())

        def _add(u: str) -> None:
            uu = _prefer_legacy_360doc_http((u or "").strip())
            if not uu:
                return
            # 带签名 URL 若已过期或将过期，直接跳过，避免优先命中旧签名。
            try:
                pu = urlparse(uu)
                if pu.query and "signature=" in pu.query.lower():
                    qs = dict(parse_qsl(pu.query, keep_blank_values=True))
                    exp_raw = str(qs.get("Expires", "")).strip()
                    if exp_raw.isdigit() and int(exp_raw) <= now_ts + 15:
                        return
            except Exception:
                pass
            if uu in cands:
                return
            cands.append(uu)

        _add(url)
        _add(_https_variant(url))
        _add(_strip_url_query(url))
        _add(_https_variant(_strip_url_query(url)))

        tg = url_primary_tag.get(url)
        raw_primary_host = ""
        if isinstance(tg, Tag):
            raw_src = str(tg.get("src", "")).strip()
            src_abs = _prefer_working_360doc_image_host(normalize_url(raw_src, source_url))
            if is_localizable_url(src_abs):
                try:
                    raw_primary_host = (urlparse(src_abs).netloc or "").lower()
                except Exception:
                    raw_primary_host = ""
                _add(src_abs)
                _add(_https_variant(src_abs))
                _add(_strip_url_query(src_abs))
                _add(_https_variant(_strip_url_query(src_abs)))

        fm, pm = _get_signed_maps()
        if url in fm:
            _add(fm[url])
            _add(_https_variant(fm[url]))
        nq = _strip_url_query(url)
        if nq in fm:
            _add(fm[nq])
            _add(_https_variant(fm[nq]))
        pkey = _url_path_key(url)
        if pkey and pkey in pm:
            _add(pm[pkey])
            _add(_https_variant(pm[pkey]))

        # Frontend fallback path: call imgurl.ashx?op=changeurl to request a fresh backend-signed URL.
        with cache_lock:
            cached_changed = changed_url_cache.get(url)
        if not cached_changed:
            pkey2 = _url_path_key(url)
            if pkey2:
                with cache_lock:
                    cached_changed = changed_path_map.get(pkey2, "")
        if not cached_changed:
            try:
                by_input, by_path = _request_changeurl_signed_images(
                    session, source_url, [url]
                )
                fresh = by_input.get(url, "")
                if not fresh:
                    fresh = by_path.get(_url_path_key(url), "")
                with cache_lock:
                    if fresh:
                        changed_url_cache[url] = fresh
                        p = _url_path_key(fresh)
                        if p:
                            changed_path_map[p] = fresh
                cached_changed = fresh
            except Exception:
                cached_changed = ""
        if cached_changed:
            _add(cached_changed)
            _add(_https_variant(cached_changed))
            # 域名容灾：同一签名 URL 在单节点 502 时，改写到同族 host 再试。
            try:
                ch_host = (urlparse(cached_changed).netloc or "").lower()
            except Exception:
                ch_host = ""
            fallback_hosts: list[str] = []
            if raw_primary_host:
                fallback_hosts.extend(_legacy_360doc_host_family(raw_primary_host))
            fallback_hosts.extend(_legacy_360doc_host_family(ch_host))
            for fh in fallback_hosts:
                if fh == ch_host:
                    continue
                rw = _rewrite_url_host(cached_changed, fh)
                _add(rw)
                _add(_https_variant(rw))

        # 额外容错：若已有 Signature，按当前时间戳和分类 artnum(domain) 组一个变体再试一次。
        # 说明：不改变 Signature 本体，仅做参数层面的轻量补参重试。
        for base in list(cands):
            try:
                p = urlparse(base)
                if not p.query or "signature=" not in p.query.lower():
                    continue
                qs = dict(parse_qsl(p.query, keep_blank_values=True))
                qs["Expires"] = str(int(time.time()))
                if domain_hint:
                    qs["domain"] = domain_hint
                q2 = urlencode(qs, doseq=True)
                _add(urlunparse((p.scheme, p.netloc, p.path, p.params, q2, p.fragment)))
            except Exception:
                continue
        return cands

    def _fetch_one(url: str) -> tuple[str, str, bytes | None, str | None, Exception | None]:
        # Return tuple: (url, status, data, ext, err), where status is ok|not_found|failed.
        time.sleep(random.uniform(*RESOURCE_START_JITTER_SEC))
        last_exc: Exception | None = None
        saw_not_found = False
        saw_non_not_found = False
        candidates = list(_candidate_urls_for_primary(url))
        tried: set[str] = set()
        attempts = 0
        refresh_retries = 0
        while candidates and attempts < RESOURCE_MAX_ATTEMPTS_PER_URL:
            cu = (candidates.pop(0) or "").strip()
            if not cu or cu in tried:
                continue
            tried.add(cu)
            attempts += 1
            try:
                resp = request_with_retry(
                    session,
                    cu,
                    headers={
                        "Referer": source_url,
                        "Accept": (
                            "image/avif,image/webp,image/apng,image/*,*/*;q=0.8"
                        ),
                        "Accept-Language": "zh-CN,zh;q=0.9,en;q=0.8",
                        "Sec-Fetch-Dest": "image",
                        "Sec-Fetch-Mode": "no-cors",
                        "Sec-Fetch-Site": "cross-site",
                        "User-Agent": session.headers.get("User-Agent", "Mozilla/5.0"),
                    },
                    timeout=RESOURCE_REQUEST_TIMEOUT,
                    retries=RESOURCE_REQUEST_RETRIES,
                    use_session_cookies=("signature=" not in cu.lower()),
                    bypass_env_proxy=True,
                )
                _, fallback_ext = url_meta[url]
                ext = suffix_from_url(cu, fallback_ext)
                magic_ext = _suffix_from_magic(resp.content)
                if magic_ext:
                    ext = magic_ext
                if ext in (".bin", ".html"):
                    ext = _suffix_from_content_type(
                        resp.headers.get("content-type", ""),
                        ext,
                    )
                return (url, "ok", resp.content, ext, None)
            except CleanRateLimitError:
                raise
            except ResourceExpiredError as exc:
                last_exc = exc
                saw_non_not_found = True
                # 惰性刷新：签名过期时实时向 changeurl 再要一次新签名 URL。
                if refresh_retries >= RESOURCE_MAX_REFRESH_RETRIES:
                    continue
                try:
                    by_input, by_path = _request_changeurl_signed_images(
                        session, source_url, [url]
                    )
                    fresh = by_input.get(url, "")
                    if not fresh:
                        fresh = by_path.get(_url_path_key(url), "")
                    if fresh:
                        with cache_lock:
                            changed_url_cache[url] = fresh
                            p = _url_path_key(fresh)
                            if p:
                                changed_path_map[p] = fresh
                        if fresh not in tried:
                            candidates.insert(0, fresh)
                            refresh_retries += 1
                except Exception:
                    pass
                continue
            except ResourceGatewayError as exc:
                last_exc = exc
                saw_non_not_found = True
                # 网关抖动（502/503/504）时同样刷新一条新签名 URL 再试。
                if refresh_retries >= RESOURCE_MAX_REFRESH_RETRIES:
                    continue
                try:
                    by_input, by_path = _request_changeurl_signed_images(
                        session, source_url, [url]
                    )
                    fresh = by_input.get(url, "")
                    if not fresh:
                        fresh = by_path.get(_url_path_key(url), "")
                    if fresh:
                        with cache_lock:
                            changed_url_cache[url] = fresh
                            p = _url_path_key(fresh)
                            if p:
                                changed_path_map[p] = fresh
                        if fresh not in tried:
                            candidates.insert(0, fresh)
                            refresh_retries += 1
                except Exception:
                    pass
                continue
            except ResourceNotFoundError as exc:
                saw_not_found = True
                last_exc = exc
                # 404/NotFound 时也尝试刷新一条签名 URL 再试，避免直链 check* 404 误判。
                if refresh_retries >= RESOURCE_MAX_REFRESH_RETRIES:
                    continue
                try:
                    by_input, by_path = _request_changeurl_signed_images(
                        session, source_url, [url]
                    )
                    fresh = by_input.get(url, "")
                    if not fresh:
                        fresh = by_path.get(_url_path_key(url), "")
                    if fresh:
                        with cache_lock:
                            changed_url_cache[url] = fresh
                            p = _url_path_key(fresh)
                            if p:
                                changed_path_map[p] = fresh
                        if fresh not in tried:
                            candidates.insert(0, fresh)
                            refresh_retries += 1
                except Exception:
                    pass
                continue
            except Exception as exc:
                if isinstance(exc, requests.exceptions.SSLError) and cu.startswith("https://"):
                    fallback_http = "http://" + cu[len("https://") :]
                    fallback_http = _prefer_legacy_360doc_http(fallback_http)
                    if fallback_http not in tried:
                        candidates.insert(0, fallback_http)
                last_exc = exc
                saw_non_not_found = True
                continue
        if saw_not_found and not saw_non_not_found and last_exc is not None:
            return (url, "not_found", None, None, last_exc)
        if attempts >= RESOURCE_MAX_ATTEMPTS_PER_URL and last_exc is not None:
            return (
                url,
                "failed",
                None,
                None,
                RuntimeError(
                    f"resource retry exhausted attempts={attempts} refresh={refresh_retries} last={last_exc}"
                ),
            )
        return (url, "failed", None, None, last_exc or RuntimeError("resource download failed"))

    workers = max(1, min(RESOURCE_DOWNLOAD_MAX_WORKERS, len(unique_urls)))
    log_info(
        f"资源下载开始 art={article_id} total={len(unique_urls)} workers={workers}"
    )
    results: dict[str, tuple[str, bytes | None, str | None, Exception | None]] = {}
    ex = concurrent.futures.ThreadPoolExecutor(max_workers=workers)
    try:
        fut2url = {ex.submit(_fetch_one, u): u for u in unique_urls}
        pending: set[concurrent.futures.Future] = set(fut2url.keys())
        done_count = 0
        last_hb = time.time()
        while pending:
            done_now, pending = concurrent.futures.wait(
                pending,
                timeout=1.0,
                return_when=concurrent.futures.FIRST_COMPLETED,
            )
            for fut in done_now:
                u = fut2url[fut]
                try:
                    _, st, data, ext, err = fut.result()
                except CleanRateLimitError:
                    raise
                results[u] = (st, data, ext, err)
                done_count += 1
                if done_count == 1 or done_count % 10 == 0 or done_count == len(unique_urls):
                    log_info(
                        f"资源下载进度 art={article_id} {done_count}/{len(unique_urls)}"
                    )
            now = time.time()
            if now - last_hb >= RESOURCE_PROGRESS_HEARTBEAT_SEC and pending:
                log_info(
                    f"资源下载心跳 art={article_id} done={done_count} pending={len(pending)}"
                )
                last_hb = now
    except KeyboardInterrupt:
        log_warn(f"资源下载被中断 art={article_id}，正在取消剩余任务...")
        ex.shutdown(wait=False, cancel_futures=True)
        raise
    finally:
        ex.shutdown(wait=False, cancel_futures=True)

    for abs_url in unique_urls:
        st, data, ext, err = results.get(abs_url, ("failed", None, None, RuntimeError("missing result")))
        seq, fallback_ext = url_meta[abs_url]
        if st == "ok" and data is not None and ext is not None:
            if not res_dir.exists():
                res_dir.mkdir(parents=True, exist_ok=True)
            local_name = sanitize_name(
                f"res_{seq}{ext}", f"res_{seq}{fallback_ext}"
            )
            local_path = res_dir / local_name
            primary_tag = url_primary_tag.get(abs_url)
            if (
                isinstance(primary_tag, Tag)
                and (primary_tag.name or "").lower() == "link"
                and local_name.lower().endswith(".css")
            ):
                data = _rewrite_css_and_localize_deps(data, abs_url, local_name)
            local_path.write_bytes(data)
            rel_ref = f"{res_dir.name}/{local_name}"
            url_to_local[abs_url] = rel_ref
            downloaded += 1
        elif st == "not_found":
            append_resource_not_found_warning_line(
                f"article_id={article_id}\tarticle={article_title}\tdir={article_dir_name}"
                f"\tresource={abs_url}\tnot_found=1"
            )
        else:
            exc = err or RuntimeError("resource download failed")
            log_warn(f"资源下载失败 {abs_url} err={exc}")
            failed_urls.append(abs_url)
            append_clean_resource_failure_line(
                article_id=article_id,
                article_title=article_title,
                article_dir_name=article_dir_name,
                resource_url=abs_url,
                error=exc,
            )
        time.sleep(random.uniform(*RESOURCE_REQUEST_SLEEP_SEC))

    for tag, _, write_attr, abs_url in plan:
        ref = url_to_local.get(abs_url)
        if ref:
            tag[write_attr] = ref
    _heal_imgs_missing_src_from_parent_anchor(clean_soup)
    return ResourceLocalizationResult(downloaded=downloaded, failed_urls=failed_urls)


def _heal_imgs_missing_src_from_parent_anchor(soup: BeautifulSoup) -> None:
    # When img src is missing, fallback to parent <a href> local path and fill src.
    root = soup.select_one("#content")
    if root is None:
        return
    for img in root.find_all("img"):
        if not isinstance(img, Tag):
            continue
        raw = str(img.get("src", "")).strip()
        if raw and not raw.startswith("data:"):
            continue
        p = img.parent
        if not isinstance(p, Tag) or (p.name or "").lower() != "a":
            continue
        h = str(p.get("href", "")).strip()
        hl = h.lower()
        if h and any(hl.endswith(ext) for ext in _LOCAL_IMAGE_HREF_EXTS):
            img["src"] = h


def extract_article_id(path: Path) -> str:
    match = re.match(r"(\d+)-", path.name)
    if match:
        return match.group(1)
    return "unknown"


def guess_source_url(path: Path) -> str:
    artid = extract_article_id(path)
    if artid != "unknown":
        return f"{BASE_URL}/showweb/0/0/{artid}.aspx"
    return BASE_URL


def _parse_css_color(s: str) -> RGBColor | None:
    s = s.strip()
    m = re.search(r"#([0-9a-fA-F]{3}|[0-9a-fA-F]{6})\b", s)
    if m:
        h = m.group(1)
        if len(h) == 3:
            r, g, b = (int(h[i] + h[i], 16) for i in range(3))
        else:
            r, g, b = int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16)
        return RGBColor(r, g, b)
    m = re.search(
        r"rgba?\(\s*(\d+)\s*,\s*(\d+)\s*,\s*(\d+)", s, re.I
    )
    if m:
        return RGBColor(int(m.group(1)), int(m.group(2)), int(m.group(3)))
    return None


def _parse_css_font_size_pt(style: str) -> float | None:
    m = re.search(r"font-size\s*:\s*([\d.]+)\s*(px|pt)", style, re.I)
    if not m:
        return None
    val = float(m.group(1))
    unit = m.group(2).lower()
    if unit == "pt":
        return val
    if unit == "px":
        return val * 0.75
    return None


def _parse_css_font_size_pt_value(value: str) -> float | None:
    m = re.search(r"([\d.]+)\s*(px|pt)\b", value or "", re.I)
    if not m:
        return None
    val = float(m.group(1))
    unit = m.group(2).lower()
    if unit == "pt":
        return val
    if unit == "px":
        return val * 0.75
    return None


def _css_decl_map(decl_text: str) -> dict[str, str]:
    out: dict[str, str] = {}
    for part in (decl_text or "").split(";"):
        if ":" not in part:
            continue
        k, v = part.split(":", 1)
        kk = k.strip().lower()
        vv = v.strip()
        if kk:
            out[kk] = vv
    return out


def _clean_font_family_name(raw: str) -> str | None:
    if not raw:
        return None
    first = raw.split(",", 1)[0].strip().strip("\"'")
    if not first:
        return None
    if re.fullmatch(r"ff\d+", first, re.I):
        return None
    return first


def _apply_css_decl_to_ctx(base: "_RunCtx", decl: dict[str, str]) -> "_RunCtx":
    ctx = base
    color = _parse_css_color(decl.get("color", "")) if "color" in decl else None
    if color is not None:
        ctx = ctx._replace(color=color)
    if "font-size" in decl:
        pt = _parse_css_font_size_pt_value(decl.get("font-size", ""))
        if pt is not None:
            ctx = ctx._replace(font_pt=pt)
    if "font-family" in decl:
        fam = _clean_font_family_name(decl.get("font-family", ""))
        if fam:
            ctx = ctx._replace(font_name=fam)
    if "font-style" in decl and decl.get("font-style", "").strip().lower() == "italic":
        ctx = ctx._replace(italic=True)
    if "font-weight" in decl:
        fw = decl.get("font-weight", "").strip().lower()
        if fw == "bold":
            ctx = ctx._replace(bold=True)
        else:
            m_num = re.search(r"\d+", fw)
            if m_num and int(m_num.group(0)) >= 600:
                ctx = ctx._replace(bold=True)
    td = decl.get("text-decoration", "").strip().lower()
    if td:
        if "underline" in td:
            ctx = ctx._replace(underline=True)
        if "line-through" in td:
            ctx = ctx._replace(strike=True)
    return ctx


def _get_docx_class_ctx_map() -> dict[str, dict[str, str]]:
    raw = getattr(_DOCX_TLS, "class_ctx_map", None)
    if isinstance(raw, dict):
        return raw
    return {}


def _parse_css_len_pt(value: str) -> float | None:
    m = re.search(r"(-?[\d.]+)\s*(px|pt)\b", value or "", re.I)
    if not m:
        return None
    val = float(m.group(1))
    unit = m.group(2).lower()
    if unit == "px":
        return val * 0.75
    if unit == "pt":
        return val
    return None


def _node_css_decl(tag: Tag) -> dict[str, str]:
    out: dict[str, str] = {}
    class_map = _get_docx_class_ctx_map()
    for cls in _tag_class_list(tag):
        decl = class_map.get(cls)
        if isinstance(decl, dict):
            out.update(decl)
    st = str(tag.get("style") or "")
    if st:
        out.update(_css_decl_map(st))
    return out


def _apply_node_paragraph_css(p, tag: Tag) -> None:
    decl = _node_css_decl(tag)
    ta = (decl.get("text-align") or "").strip().lower()
    if ta == "center":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif ta == "right":
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    elif ta == "justify":
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    elif ta == "left":
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    pf = p.paragraph_format
    left_pt = None
    for k in ("margin-left", "padding-left"):
        v = decl.get(k)
        if not v:
            continue
        vv = _parse_css_len_pt(v)
        if vv is None:
            continue
        if left_pt is None or vv > left_pt:
            left_pt = vv
    if left_pt is not None and left_pt > 0:
        pf.left_indent = Pt(left_pt)

    ti = decl.get("text-indent")
    if ti:
        ti_pt = _parse_css_len_pt(ti)
        if ti_pt is not None:
            pf.first_line_indent = Pt(ti_pt)


def _is_preview_image_tag(img: Tag) -> bool:
    if not isinstance(img, Tag):
        return False
    cls = set(_tag_class_list(img))
    if "bi" in cls:
        return True
    alt = str(img.get("alt", "")).lower()
    if alt.startswith("pdf-page-") or alt.startswith("ppt-page-"):
        return True
    for p in img.parents:
        if not isinstance(p, Tag):
            continue
        c = set(_tag_class_list(p))
        if {
            "word-preview-page",
            "pdf-preview-page",
            "ppt-preview-page",
            "word-document-preview",
            "pdf-document-preview",
            "ppt-document-preview",
        } & c:
            return True
    return False


def _nearest_word_preview_page(tag: Tag) -> Tag | None:
    for p in tag.parents:
        if not isinstance(p, Tag):
            continue
        cls = set(_tag_class_list(p))
        if "word-preview-page" in cls:
            return p
    return None


def _word_preview_page_has_text(page: Tag) -> bool:
    for t in page.find_all(class_="t"):
        if not isinstance(t, Tag):
            continue
        if t.get_text("", strip=True):
            return True
    return False


def _is_word_preview_background_image(img: Tag) -> bool:
    cls = set(_tag_class_list(img))
    if "bi" not in cls:
        return False
    # Typical page background sprite in preview template.
    if {"x0", "y0", "w1", "h1"}.issubset(cls):
        return True
    return False


def _should_skip_preview_image_for_docx(img: Tag) -> bool:
    if not isinstance(img, Tag):
        return False
    if not _is_word_preview_background_image(img):
        return False
    page = _nearest_word_preview_page(img)
    if page is None:
        return False
    return _word_preview_page_has_text(page)


def _image_width_for_docx(img: Tag, default_width: Inches = Inches(5.5)):
    decl = _node_css_decl(img)
    w = decl.get("width", "")
    pt = _parse_css_len_pt(w)
    if pt is None:
        return default_width
    inches = pt / 72.0
    # Clamp to a practical page width range.
    if inches < 0.5:
        inches = 0.5
    if inches > 7.0:
        inches = 7.0
    return Inches(inches)


def _set_run_east_asia(run, font_name: str) -> None:
    run.font.name = font_name
    r = run._element.rPr
    if r is not None and r.rFonts is not None:
        r.rFonts.set(qn("w:eastAsia"), font_name)


def _init_doc_typography(doc: Document) -> None:
    normal = doc.styles["Normal"]
    npf = normal.paragraph_format
    npf.space_before = Pt(0)
    npf.space_after = Pt(0)
    npf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    npf.line_spacing = FIXED_LINE_SPACING_PT


def _apply_body_paragraph_format(p) -> None:
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    pf.line_spacing = FIXED_LINE_SPACING_PT


def _paragraph_has_visible_content(p: Paragraph) -> bool:
    raw = (p.text or "").replace("\xa0", " ").replace("\u200b", "")
    if raw:
        return True
    for run in p.runs:
        el = run._element
        if el.find(qn("w:drawing")) is not None:
            return True
        if el.find(qn("w:pict")) is not None:
            return True
    return False


def _remove_empty_paragraphs(doc: Document) -> None:
    body_el = doc._body._body
    for child in list(body_el):
        if child.tag != qn("w:p"):
            continue
        para = Paragraph(child, doc._body)
        if not _paragraph_has_visible_content(para):
            body_el.remove(child)


def _apply_paragraph_format_in_table(tbl) -> None:
    for row in tbl.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                _apply_body_paragraph_format(p)
            for nested in cell.tables:
                _apply_paragraph_format_in_table(nested)


def _apply_body_paragraph_format_to_all(doc: Document) -> None:
    for p in doc.paragraphs:
        _apply_body_paragraph_format(p)
    for tbl in doc.tables:
        _apply_paragraph_format_in_table(tbl)


class _RunCtx(NamedTuple):
    bold: bool = False
    italic: bool = False
    underline: bool = False
    strike: bool = False
    subscript: bool = False
    superscript: bool = False
    color: RGBColor | None = None
    font_pt: float | None = None
    font_name: str | None = None


_LINK_BLUE = RGBColor(0x05, 0x63, 0xC1)
_DOCX_TLS = threading.local()


def _rgb_word_hex(rgb: RGBColor) -> str:
    return "%02X%02X%02X" % (rgb[0], rgb[1], rgb[2])


def _apply_ctx_to_run(run, ctx: _RunCtx, default_pt: Pt) -> None:
    fn = ctx.font_name or "绛夌嚎"
    _set_run_east_asia(run, fn)
    if ctx.font_pt is not None:
        run.font.size = Pt(ctx.font_pt)
    else:
        run.font.size = default_pt
    run.bold = ctx.bold
    run.italic = ctx.italic
    run.font.underline = (
        WD_UNDERLINE.SINGLE if ctx.underline else WD_UNDERLINE.NONE
    )
    run.font.strike = ctx.strike
    run.font.subscript = bool(ctx.subscript)
    run.font.superscript = bool(ctx.superscript)
    if ctx.color is not None:
        run.font.color.rgb = ctx.color


def _span_ctx_from_tag(tag: Tag, base: _RunCtx) -> _RunCtx:
    ctx = base
    class_map = _get_docx_class_ctx_map()
    for cls in _tag_class_list(tag):
        decl = class_map.get(cls)
        if isinstance(decl, dict):
            ctx = _apply_css_decl_to_ctx(ctx, decl)

    st = str(tag.get("style") or "")
    if st:
        ctx = _apply_css_decl_to_ctx(ctx, _css_decl_map(st))

    if tag.has_attr("color"):
        co = str(tag.get("color", ""))
        c = _parse_css_color(co)
        if c is not None:
            ctx = ctx._replace(color=c)
    return ctx


def _append_ox_text_run(
    ox_parent,
    text: str,
    ctx: _RunCtx,
    default_pt: Pt,
    *,
    link_style: bool,
) -> None:
    r = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    if ctx.bold:
        r_pr.append(OxmlElement("w:b"))
    if ctx.italic:
        r_pr.append(OxmlElement("w:i"))
    color = ctx.color
    if color is None and link_style:
        color = _LINK_BLUE
    if color is not None:
        co = OxmlElement("w:color")
        co.set(qn("w:val"), _rgb_word_hex(color))
        r_pr.append(co)
    if ctx.underline or link_style:
        u = OxmlElement("w:u")
        u.set(qn("w:val"), "single")
        r_pr.append(u)
    if ctx.strike:
        r_pr.append(OxmlElement("w:strike"))
    if ctx.subscript:
        va = OxmlElement("w:vertAlign")
        va.set(qn("w:val"), "subscript")
        r_pr.append(va)
    if ctx.superscript:
        va = OxmlElement("w:vertAlign")
        va.set(qn("w:val"), "superscript")
        r_pr.append(va)
    pt_val = ctx.font_pt if ctx.font_pt is not None else float(default_pt.pt)
    half = str(int(round(pt_val * 2)))
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), half)
    r_pr.append(sz)
    sz_cs = OxmlElement("w:szCs")
    sz_cs.set(qn("w:val"), half)
    r_pr.append(sz_cs)
    r_fonts = OxmlElement("w:rFonts")
    fn = ctx.font_name or "绛夌嚎"
    r_fonts.set(qn("w:ascii"), fn)
    r_fonts.set(qn("w:hAnsi"), fn)
    r_fonts.set(qn("w:eastAsia"), fn)
    r_pr.append(r_fonts)
    r.append(r_pr)
    t_el = OxmlElement("w:t")
    t_el.text = text
    t_el.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    r.append(t_el)
    ox_parent.append(r)


def _walk_inline_to_hyperlink_ox(
    h_el,
    node,
    ctx: _RunCtx,
    default_pt: Pt,
    *,
    link_style: bool,
    br_state: list[bool] | None = None,
) -> None:
    if br_state is None:
        br_state = [False]
    if isinstance(node, NavigableString):
        t = _normalize_navigable_text_for_docx(str(node))
        if not t:
            return
        if not t.strip():
            if "\n" in t or "\r" in t:
                return
            br_state[0] = False
            _append_ox_text_run(h_el, t, ctx, default_pt, link_style=link_style)
            return
        br_state[0] = False
        _append_ox_text_run(h_el, t, ctx, default_pt, link_style=link_style)
        return
    if not isinstance(node, Tag):
        return
    name = (node.name or "").lower()
    node_ctx = _span_ctx_from_tag(node, ctx)
    if name == "br":
        if br_state[0]:
            return
        h_el.append(OxmlElement("w:br"))
        br_state[0] = True
        return
    if name in ("strong", "b"):
        for ch in node.children:
            _walk_inline_to_hyperlink_ox(
                h_el,
                ch,
                node_ctx._replace(bold=True),
                default_pt,
                link_style=link_style,
                br_state=br_state,
            )
        return
    if name in ("em", "i"):
        for ch in node.children:
            _walk_inline_to_hyperlink_ox(
                h_el,
                ch,
                node_ctx._replace(italic=True),
                default_pt,
                link_style=link_style,
                br_state=br_state,
            )
        return
    if name == "u":
        for ch in node.children:
            _walk_inline_to_hyperlink_ox(
                h_el,
                ch,
                node_ctx._replace(underline=True),
                default_pt,
                link_style=link_style,
                br_state=br_state,
            )
        return
    if name in ("s", "strike", "del"):
        for ch in node.children:
            _walk_inline_to_hyperlink_ox(
                h_el,
                ch,
                node_ctx._replace(strike=True),
                default_pt,
                link_style=link_style,
                br_state=br_state,
            )
        return
    if name == "sub":
        for ch in node.children:
            _walk_inline_to_hyperlink_ox(
                h_el,
                ch,
                node_ctx._replace(subscript=True, superscript=False),
                default_pt,
                link_style=link_style,
                br_state=br_state,
            )
        return
    if name == "sup":
        for ch in node.children:
            _walk_inline_to_hyperlink_ox(
                h_el,
                ch,
                node_ctx._replace(superscript=True, subscript=False),
                default_pt,
                link_style=link_style,
                br_state=br_state,
            )
        return
    if name in ("span", "font"):
        for ch in node.children:
            _walk_inline_to_hyperlink_ox(
                h_el,
                ch,
                node_ctx,
                default_pt,
                link_style=link_style,
                br_state=br_state,
            )
        return
    if name in ("code",):
        for ch in node.children:
            _walk_inline_to_hyperlink_ox(
                h_el,
                ch,
                node_ctx._replace(font_name="Consolas"),
                default_pt,
                link_style=link_style,
                br_state=br_state,
            )
        return
    for ch in node.children:
        _walk_inline_to_hyperlink_ox(
            h_el, ch, node_ctx, default_pt, link_style=link_style, br_state=br_state
        )


# Base z-order for floating anchors, preventing overlap from repeated relative-height values.
_ANCHOR_RELATIVE_HEIGHT_NEXT = 251_658_240


def _next_anchor_relative_height() -> int:
    global _ANCHOR_RELATIVE_HEIGHT_NEXT
    _ANCHOR_RELATIVE_HEIGHT_NEXT += 1
    return _ANCHOR_RELATIVE_HEIGHT_NEXT


def _strip_line_breaks_after_drawing_in_run(run) -> None:
    # Remove soft breaks right after drawings to avoid extra blank lines with top/bottom wrapping.
    r_el = run._element
    seen_drawing = False
    for child in list(r_el):
        if child.tag == qn("w:drawing"):
            seen_drawing = True
            continue
        if seen_drawing and child.tag == qn("w:br"):
            r_el.remove(child)


def _convert_run_inline_picture_to_top_bottom_wrap(run) -> None:
    # Convert wp:inline from add_picture() into wp:anchor with top-and-bottom text wrapping.
    r_el = run._element
    drawing = r_el.find(qn("w:drawing"))
    if drawing is None:
        return
    inline = drawing.find(qn("wp:inline"))
    if inline is None:
        return
    extent = inline.find(qn("wp:extent"))
    effect = inline.find(qn("wp:effectExtent"))
    doc_pr = inline.find(qn("wp:docPr"))
    cnv = inline.find(qn("wp:cNvGraphicFramePr"))
    graphic = inline.find(qn("a:graphic"))
    if extent is None or graphic is None or doc_pr is None:
        return

    z = _next_anchor_relative_height()
    extent_c = copy.deepcopy(extent)
    if effect is not None:
        effect_c = copy.deepcopy(effect)
    else:
        effect_c = parse_xml(
            '<wp:effectExtent %s l="0" t="0" r="0" b="0"/>' % nsdecls("wp")
        )
    doc_pr_c = copy.deepcopy(doc_pr)
    cnv_c = copy.deepcopy(cnv) if cnv is not None else None
    graphic_c = copy.deepcopy(graphic)

    decl = nsdecls("wp", "a", "pic", "r")
    anchor = parse_xml(
        "<wp:anchor %s distT=\"0\" distB=\"0\" distL=\"0\" distR=\"0\" "
        'simplePos="0" relativeHeight="%d" behindDoc="0" locked="0" '
        'layoutInCell="1" allowOverlap="0"></wp:anchor>' % (decl, z)
    )
    sp = parse_xml('<wp:simplePos %s x="0" y="0"/>' % nsdecls("wp"))
    pos_h = parse_xml(
        "<wp:positionH %s relativeFrom=\"column\">"
        "<wp:align>center</wp:align></wp:positionH>" % nsdecls("wp")
    )
    pos_v = parse_xml(
        "<wp:positionV %s relativeFrom=\"paragraph\">"
        "<wp:posOffset>0</wp:posOffset></wp:positionV>" % nsdecls("wp")
    )
    wrap_tb = parse_xml("<wp:wrapTopAndBottom %s/>" % nsdecls("wp"))

    anchor.append(sp)
    anchor.append(pos_h)
    anchor.append(pos_v)
    anchor.append(extent_c)
    anchor.append(effect_c)
    anchor.append(wrap_tb)
    anchor.append(doc_pr_c)
    if cnv_c is not None:
        anchor.append(cnv_c)
    anchor.append(graphic_c)

    drawing.clear()
    drawing.append(anchor)


def _add_picture_top_bottom_wrap(paragraph, image_path: str, width) -> None:
    run = paragraph.add_run()
    try:
        run.add_picture(str(image_path), width=width)
    except (UnrecognizedImageError, OSError, ValueError):
        # Fallback for WEBP/CMYK/legacy formats that python-docx cannot parse directly.
        with Image.open(str(image_path)) as im:
            try:
                im.seek(0)
            except Exception:
                pass
            if im.mode not in ("RGB", "RGBA", "L", "LA"):
                im = im.convert("RGBA")
            buf = io.BytesIO()
            im.save(buf, format="PNG")
            buf.seek(0)
            run.add_picture(buf, width=width)
    _convert_run_inline_picture_to_top_bottom_wrap(run)
    _strip_line_breaks_after_drawing_in_run(run)


def _exc_brief(exc: Exception) -> str:
    msg = str(exc).strip()
    if msg:
        return f"{type(exc).__name__}: {msg}"
    return type(exc).__name__


def _paragraph_add_external_hyperlink(
    paragraph,
    url: str,
    anchor: Tag,
    ctx: _RunCtx,
    default_pt: Pt,
    *,
    br_state: list[bool] | None = None,
) -> None:
    if br_state is None:
        br_state = [False]
    part = paragraph.part
    r_id = part.relate_to(url, RT.HYPERLINK, is_external=True)
    h_el = OxmlElement("w:hyperlink")
    h_el.set(qn("r:id"), r_id)
    paragraph._p.append(h_el)
    _walk_inline_to_hyperlink_ox(
        h_el, anchor, ctx, default_pt, link_style=True, br_state=br_state
    )


def _add_inline_image_to_paragraph(
    paragraph,
    img: Tag,
    base_dir: Path,
    *,
    article_clean_html: Path | None = None,
    br_state: list[bool] | None = None,
) -> None:
    if _should_skip_preview_image_for_docx(img):
        return
    if br_state is not None:
        br_state[0] = False
    src = _img_src_for_local_resolve(img)
    local = _resolve_local_media_path(
        src, base_dir, article_clean_html=article_clean_html
    )
    if local and local.is_file():
        try:
            _add_picture_top_bottom_wrap(
                paragraph, str(local), _image_width_for_docx(img)
            )
        except Exception as exc:
            log_warn(f"插入图片失败 {local}: {_exc_brief(exc)}")
            r2 = paragraph.add_run(f"[鍥剧墖 {src}]")
            _set_run_east_asia(r2, "绛夌嚎")
            r2.font.size = DEFAULT_BODY_PT
    else:
        r2 = paragraph.add_run(f"[鍥剧墖 {src}]")
        _set_run_east_asia(r2, "绛夌嚎")


def _tag_class_list(tag: Tag) -> list[str]:
    c = tag.get("class") or []
    if isinstance(c, str):
        return c.split()
    return list(c)


def _normalize_navigable_text_for_docx(raw: str) -> str:
    # Normalize common NBSP/space padding patterns to avoid one-character-per-line layout artifacts.
    if not raw:
        return raw
    t = raw.replace("\u00a0", " ").replace("\u2003", " ").replace("\u2002", " ")
    t = re.sub(r"[ \t]+", " ", t)
    t = re.sub(r" +\n", "\n", t)
    t = re.sub(r"\n{3,}", "\n\n", t)
    return t


def _walk_inline_to_paragraph(
    paragraph,
    node,
    ctx: _RunCtx,
    default_pt: Pt,
    *,
    base_dir: Path | None = None,
    article_clean_html: Path | None = None,
    br_state: list[bool] | None = None,
    local_media_href_dedupe: set[str] | None = None,
) -> None:
    if br_state is None:
        br_state = [False]

    def _w(ch, cctx: _RunCtx) -> None:
        _walk_inline_to_paragraph(
            paragraph,
            ch,
            cctx,
            default_pt,
            base_dir=base_dir,
            article_clean_html=article_clean_html,
            br_state=br_state,
            local_media_href_dedupe=local_media_href_dedupe,
        )

    if isinstance(node, NavigableString):
        t = _normalize_navigable_text_for_docx(str(node))
        if not t:
            return
        if not t.strip():
            if "\n" in t or "\r" in t:
                return
            br_state[0] = False
            run = paragraph.add_run(t)
            _apply_ctx_to_run(run, ctx, default_pt)
            return
        br_state[0] = False
        run = paragraph.add_run(t)
        _apply_ctx_to_run(run, ctx, default_pt)
        return
    if not isinstance(node, Tag):
        return
    name = (node.name or "").lower()
    node_ctx = _span_ctx_from_tag(node, ctx)
    if name == "br":
        if br_state[0]:
            return
        paragraph.add_run().add_break()
        br_state[0] = True
        return
    if name == "img":
        if base_dir is not None:
            src = _img_src_for_local_resolve(node)
            key = src.replace("\\", "/") if src else ""
            if key and local_media_href_dedupe is not None:
                if key in local_media_href_dedupe:
                    return
                local_media_href_dedupe.add(key)
            _add_inline_image_to_paragraph(
                paragraph,
                node,
                base_dir,
                article_clean_html=article_clean_html,
                br_state=br_state,
            )
        return
    if name in ("strong", "b"):
        for ch in node.children:
            _w(ch, node_ctx._replace(bold=True))
        return
    if name in ("em", "i"):
        for ch in node.children:
            _w(ch, node_ctx._replace(italic=True))
        return
    if name == "u":
        for ch in node.children:
            _w(ch, node_ctx._replace(underline=True))
        return
    if name in ("s", "strike", "del"):
        for ch in node.children:
            _w(ch, node_ctx._replace(strike=True))
        return
    if name == "sub":
        for ch in node.children:
            _w(ch, node_ctx._replace(subscript=True, superscript=False))
        return
    if name == "sup":
        for ch in node.children:
            _w(ch, node_ctx._replace(superscript=True, subscript=False))
        return
    if name == "a":
        href = str(node.get("href") or "").strip()
        hl = href.lower()
        is_local_image_href = (
            base_dir is not None
            and href
            and not href.startswith(("#", "//"))
            and not hl.startswith(
                ("javascript:", "mailto:", "tel:", "http://", "https://", "data:")
            )
            and any(hl.endswith(ext) for ext in _LOCAL_IMAGE_HREF_EXTS)
        )
        if is_local_image_href:
            if node.find("img"):
                for ch in node.children:
                    _w(ch, node_ctx)
                return
            key = href.replace("\\", "/")
            if local_media_href_dedupe is not None and key in local_media_href_dedupe:
                return
            if local_media_href_dedupe is not None:
                local_media_href_dedupe.add(key)
            local = _resolve_local_media_path(
                href, base_dir, article_clean_html=article_clean_html
            )
            if local and local.is_file():
                if br_state is not None:
                    br_state[0] = False
                try:
                    _add_picture_top_bottom_wrap(paragraph, str(local), Inches(5.5))
                except Exception as exc:
                    log_warn(f"插入图片失败 {local}: {_exc_brief(exc)}")
            return
        if href and not href.startswith("#") and not hl.startswith(
            ("javascript:", "mailto:", "tel:")
        ):
            _paragraph_add_external_hyperlink(
                paragraph, href, node, node_ctx, default_pt, br_state=br_state
            )
        else:
            for ch in node.children:
                _w(ch, node_ctx)
        return
    if name in ("span", "font"):
        for ch in node.children:
            _w(ch, node_ctx)
        return
    if name in ("code",):
        for ch in node.children:
            _w(ch, node_ctx._replace(font_name="Consolas"))
        return
    if name in ("mark",):
        for ch in node.children:
            _w(ch, node_ctx)
        return
    if name in ("p", "div", "section", "article", "header", "footer", "center"):
        if "word-preview-page" in _tag_class_list(node):
            if base_dir is not None:
                for im in node.find_all("img"):
                    if isinstance(im, Tag):
                        _add_inline_image_to_paragraph(
                            paragraph,
                            im,
                            base_dir,
                            article_clean_html=article_clean_html,
                            br_state=br_state,
                        )
            return
        for ch in node.children:
            _w(ch, node_ctx)
        return
    for ch in node.children:
        _w(ch, node_ctx)


def _img_src_for_local_resolve(img: Tag) -> str:
    # Resolve img src with fallback to parent <a href> local path for legacy cleaned HTML.
    s = str(img.get("src", "")).strip()
    if s and not s.startswith("data:"):
        return s
    p = img.parent
    if isinstance(p, Tag) and (p.name or "").lower() == "a":
        h = str(p.get("href", "")).strip()
        hl = h.lower()
        if any(hl.endswith(ext) for ext in _LOCAL_IMAGE_HREF_EXTS):
            return h
    return s


def _res_basename_without_collision_suffix(bare: str) -> str | None:
    # res_10_4.jpg -> res_10.jpg fallback for historical collision-avoidance naming.
    m = re.match(r"^(res_\d+)_\d+(\.[^.]+)$", bare, re.I)
    if m:
        return m.group(1) + m.group(2)
    return None


def _media_rel_key_for_dedupe(tag: Tag) -> str | None:
    # De-duplicate empty-link and image-link variants in one paragraph to prevent duplicate inserts.
    im = tag.find("img")
    if isinstance(im, Tag):
        s = _img_src_for_local_resolve(im)
        if s:
            return s.replace("\\", "/")
    if (tag.name or "").lower() == "a":
        h = str(tag.get("href", "")).strip()
        if not h or h.lower().startswith(("http://", "https://", "data:")):
            return None
        hl = h.lower()
        if any(hl.endswith(ext) for ext in _LOCAL_IMAGE_HREF_EXTS):
            return h.replace("\\", "/")
    return None


def _is_block_media_tag(tag: Tag) -> bool:
    if tag.name and tag.name.lower() in ("img", "video", "audio", "embed", "iframe"):
        return True
    if tag.name and tag.name.lower() == "a":
        href = str(tag.get("href", "")).lower()
        if any(href.endswith(ext) for ext in _LOCAL_IMAGE_HREF_EXTS + (".mp4", ".mp3")):
            return True
        if tag.find("img"):
            return True
    return False


def _path_under_article_base(path: Path, base_resolved: Path) -> bool:
    # Check whether path is under base; Windows path resolution can break strict relative_to checks.
    try:
        path.resolve().relative_to(base_resolved)
        return True
    except ValueError:
        try:
            return os.path.commonpath(
                [os.path.abspath(path), os.path.abspath(base_resolved)]
            ) == os.path.abspath(base_resolved)
        except (ValueError, OSError):
            return False


def _resolve_local_media_path(
    src: str,
    base_dir: Path,
    *,
    article_clean_html: Path | None = None,
) -> Path | None:
    src = src.strip()
    if not src or src.startswith(("http://", "https://", "data:")):
        return None
    base_resolved = base_dir.resolve()
    p = (base_dir / src).resolve()
    if p.is_file() and _path_under_article_base(p, base_resolved):
        return p
    if article_clean_html is None or not article_clean_html.is_file():
        return None
    bare = Path(src.replace("\\", "/")).name
    if not bare:
        return None
    rd = res_dir_for_clean(article_clean_html)
    p2 = (rd / bare).resolve()
    if p2.is_file() and _path_under_article_base(p2, base_resolved):
        return p2
    alt = _res_basename_without_collision_suffix(bare)
    if alt:
        p3 = (rd / alt).resolve()
        if p3.is_file() and _path_under_article_base(p3, base_resolved):
            return p3
    return None


def _append_centered_picture_to_paragraph(
    p,
    im: Tag,
    base_dir: Path,
    *,
    article_clean_html: Path | None,
) -> None:
    if _should_skip_preview_image_for_docx(im):
        return
    src = _img_src_for_local_resolve(im)
    local = _resolve_local_media_path(
        src, base_dir, article_clean_html=article_clean_html
    )
    if local and local.is_file():
        try:
            _add_picture_top_bottom_wrap(p, str(local), _image_width_for_docx(im))
        except Exception as exc:
            log_warn(f"插入图片失败 {local}: {_exc_brief(exc)}")
            r2 = p.add_run(f"[鍥剧墖 {src}]")
            _set_run_east_asia(r2, "绛夌嚎")
    else:
        r2 = p.add_run(f"[鍥剧墖 {src}]")
        _set_run_east_asia(r2, "绛夌嚎")


def _fill_paragraph_with_media(
    p,
    tag: Tag,
    base_dir: Path,
    *,
    article_clean_html: Path | None = None,
    new_paragraph: Callable[[], Paragraph] | None = None,
) -> None:
    if tag.name and tag.name.lower() == "img":
        imgs = [tag] if isinstance(tag, Tag) else []
    else:
        imgs = [im for im in tag.find_all("img") if isinstance(im, Tag)]
    if not imgs:
        if (tag.name or "").lower() == "a":
            href = str(tag.get("href", "")).strip()
            hl = href.lower()
            if href and not href.startswith(
                ("http://", "https://", "data:")
            ) and any(hl.endswith(ext) for ext in _LOCAL_IMAGE_HREF_EXTS):
                local = _resolve_local_media_path(
                    href, base_dir, article_clean_html=article_clean_html
                )
                if local and local.is_file():
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    _apply_body_paragraph_format(p)
                    try:
                        _add_picture_top_bottom_wrap(p, str(local), Inches(5.5))
                    except Exception as exc:
                        log_warn(f"插入图片失败 {local}: {_exc_brief(exc)}")
                        r2 = p.add_run(f"[鍥剧墖 {href}]")
                        _set_run_east_asia(r2, "绛夌嚎")
                    return
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        _apply_body_paragraph_format(p)
        run = p.add_run(tag.get_text(strip=True) or "[濯掍綋]")
        _set_run_east_asia(run, "绛夌嚎")
        run.font.size = DEFAULT_BODY_PT
        return

    if len(imgs) > 1 and new_paragraph is not None:
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT if _is_preview_image_tag(imgs[0]) else WD_ALIGN_PARAGRAPH.CENTER
        _apply_body_paragraph_format(p)
        _append_centered_picture_to_paragraph(
            p, imgs[0], base_dir, article_clean_html=article_clean_html
        )
        for im in imgs[1:]:
            np = new_paragraph()
            np.alignment = WD_ALIGN_PARAGRAPH.LEFT if _is_preview_image_tag(im) else WD_ALIGN_PARAGRAPH.CENTER
            _apply_body_paragraph_format(np)
            _append_centered_picture_to_paragraph(
                np, im, base_dir, article_clean_html=article_clean_html
            )
        return

    p.alignment = WD_ALIGN_PARAGRAPH.LEFT if _is_preview_image_tag(imgs[0]) else WD_ALIGN_PARAGRAPH.CENTER
    _apply_body_paragraph_format(p)
    for im in imgs:
        _append_centered_picture_to_paragraph(
            p, im, base_dir, article_clean_html=article_clean_html
        )


def _add_media_paragraph(
    doc: Document,
    tag: Tag,
    base_dir: Path,
    *,
    article_clean_html: Path | None = None,
) -> None:
    if tag.name and tag.name.lower() == "img":
        img_list = [tag]
    else:
        img_list = [im for im in tag.find_all("img") if isinstance(im, Tag)]
    if not img_list:
        p = doc.add_paragraph()
        _fill_paragraph_with_media(
            p,
            tag,
            base_dir,
            article_clean_html=article_clean_html,
            new_paragraph=lambda: doc.add_paragraph(),
        )
        return
    for im in img_list:
        p = doc.add_paragraph()
        _fill_paragraph_with_media(
            p, im, base_dir, article_clean_html=article_clean_html
        )


def _paragraph_has_content(p) -> bool:
    return bool((p.text or "").strip()) or len(p.runs) > 0


def _table_row_cells(tr: Tag) -> list[Tag]:
    return [
        c
        for c in tr.children
        if isinstance(c, Tag) and (c.name or "").lower() in ("td", "th")
    ]


def _emit_table_cell_content(
    cell,
    cell_tag: Tag,
    base_dir: Path,
    *,
    article_clean_html: Path | None = None,
) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    children = list(cell_tag.children)
    for i, ch in enumerate(children):
        if isinstance(ch, NavigableString):
            if str(ch).strip():
                _walk_inline_to_paragraph(
                    p,
                    ch,
                    _RunCtx(),
                    DEFAULT_BODY_PT,
                    base_dir=base_dir,
                    article_clean_html=article_clean_html,
                )
            continue
        if not isinstance(ch, Tag):
            continue
        if _is_block_media_tag(ch):
            if _paragraph_has_content(p):
                p = cell.add_paragraph()
            _fill_paragraph_with_media(
                p,
                ch,
                base_dir,
                article_clean_html=article_clean_html,
                new_paragraph=lambda: cell.add_paragraph(),
            )
            continue
        cn = (ch.name or "").lower()
        if cn in ("p", "div", "section"):
            if i > 0 and _paragraph_has_content(p):
                p.add_run().add_break()
            for c2 in ch.children:
                _walk_inline_to_paragraph(
                    p,
                    c2,
                    _RunCtx(),
                    DEFAULT_BODY_PT,
                    base_dir=base_dir,
                    article_clean_html=article_clean_html,
                )
        else:
            _walk_inline_to_paragraph(
                p,
                ch,
                _RunCtx(),
                DEFAULT_BODY_PT,
                base_dir=base_dir,
                article_clean_html=article_clean_html,
            )


def _emit_html_table(
    doc: Document,
    table_tag: Tag,
    base_dir: Path,
    *,
    article_clean_html: Path | None = None,
) -> None:
    rows = [r for r in table_tag.find_all("tr") if isinstance(r, Tag)]
    if not rows:
        return
    ncols = max((len(_table_row_cells(tr)) for tr in rows), default=1)
    tbl = doc.add_table(rows=0, cols=ncols)
    tbl.style = "Table Grid"
    for tr in rows:
        tags = _table_row_cells(tr)
        if not tags:
            continue
        row = tbl.add_row()
        for ci in range(ncols):
            cell = row.cells[ci]
            if ci < len(tags):
                _emit_table_cell_content(
                    cell, tags[ci], base_dir, article_clean_html=article_clean_html
                )


def _list_indent_for_level(level: int) -> Pt | None:
    if level <= 0:
        return None
    return Pt(min(10 + level * 14, 120))


def _emit_ul(
    doc: Document,
    ul: Tag,
    base_dir: Path,
    *,
    list_level: int = 0,
    article_clean_html: Path | None = None,
) -> None:
    for li in ul.children:
        if not isinstance(li, Tag) or (li.name or "").lower() != "li":
            continue
        p = doc.add_paragraph(style="List Bullet")
        ind = _list_indent_for_level(list_level)
        if ind is not None:
            p.paragraph_format.left_indent = ind
        for sub in li.children:
            if isinstance(sub, NavigableString):
                if str(sub).strip():
                    _walk_inline_to_paragraph(
                        p,
                        sub,
                        _RunCtx(),
                        DEFAULT_BODY_PT,
                        base_dir=base_dir,
                        article_clean_html=article_clean_html,
                    )
                continue
            if not isinstance(sub, Tag):
                continue
            sn = (sub.name or "").lower()
            if sn == "ul":
                _emit_ul(
                    doc,
                    sub,
                    base_dir,
                    list_level=list_level + 1,
                    article_clean_html=article_clean_html,
                )
            elif sn == "ol":
                _emit_ol(
                    doc,
                    sub,
                    base_dir,
                    list_level=list_level + 1,
                    article_clean_html=article_clean_html,
                )
            elif sn in ("p", "div"):
                for c2 in sub.children:
                    _walk_inline_to_paragraph(
                        p,
                        c2,
                        _RunCtx(),
                        DEFAULT_BODY_PT,
                        base_dir=base_dir,
                        article_clean_html=article_clean_html,
                    )
            else:
                _walk_inline_to_paragraph(
                    p,
                    sub,
                    _RunCtx(),
                    DEFAULT_BODY_PT,
                    base_dir=base_dir,
                    article_clean_html=article_clean_html,
                )


def _emit_ol(
    doc: Document,
    ol: Tag,
    base_dir: Path,
    *,
    list_level: int = 0,
    article_clean_html: Path | None = None,
) -> None:
    for li in ol.children:
        if not isinstance(li, Tag) or (li.name or "").lower() != "li":
            continue
        p = doc.add_paragraph(style="List Number")
        ind = _list_indent_for_level(list_level)
        if ind is not None:
            p.paragraph_format.left_indent = ind
        for sub in li.children:
            if isinstance(sub, NavigableString):
                if str(sub).strip():
                    _walk_inline_to_paragraph(
                        p,
                        sub,
                        _RunCtx(),
                        DEFAULT_BODY_PT,
                        base_dir=base_dir,
                        article_clean_html=article_clean_html,
                    )
                continue
            if not isinstance(sub, Tag):
                continue
            sn = (sub.name or "").lower()
            if sn == "ul":
                _emit_ul(
                    doc,
                    sub,
                    base_dir,
                    list_level=list_level + 1,
                    article_clean_html=article_clean_html,
                )
            elif sn == "ol":
                _emit_ol(
                    doc,
                    sub,
                    base_dir,
                    list_level=list_level + 1,
                    article_clean_html=article_clean_html,
                )
            elif sn in ("p", "div"):
                for c2 in sub.children:
                    _walk_inline_to_paragraph(
                        p,
                        c2,
                        _RunCtx(),
                        DEFAULT_BODY_PT,
                        base_dir=base_dir,
                        article_clean_html=article_clean_html,
                    )
            else:
                _walk_inline_to_paragraph(
                    p,
                    sub,
                    _RunCtx(),
                    DEFAULT_BODY_PT,
                    base_dir=base_dir,
                    article_clean_html=article_clean_html,
                )


# Unwrap td/th/tr as containers so Word does not collapse entire body content into a single paragraph.
_BLOCK_UNWRAP = frozenset(
    {
        "article",
        "main",
        "header",
        "footer",
        "aside",
        "nav",
        "tbody",
        "thead",
        "tfoot",
        "hgroup",
        "center",
        "td",
        "th",
        "tr",
    }
)


def _emit_content_node(
    doc: Document,
    node: Tag,
    base_dir: Path,
    *,
    list_level: int = 0,
    article_clean_html: Path | None = None,
) -> None:
    name = (node.name or "").lower()
    if name in ("script", "style"):
        return
    if name == "br":
        return
    cls = _tag_class_list(node)
    base_inline_ctx = _span_ctx_from_tag(node, _RunCtx())
    if "t" in cls:
        tp = doc.add_paragraph()
        _apply_node_paragraph_css(tp, node)
        for ch in node.children:
            _walk_inline_to_paragraph(
                tp,
                ch,
                base_inline_ctx,
                DEFAULT_BODY_PT,
                base_dir=base_dir,
                article_clean_html=article_clean_html,
            )
        return
    if "word-document-preview" in cls:
        # Specialized preview template: preserve both text and media by walking pages.
        for ch in node.children:
            if isinstance(ch, NavigableString):
                if str(ch).strip():
                    bp = doc.add_paragraph()
                    _walk_inline_to_paragraph(
                        bp,
                        ch,
                        base_inline_ctx,
                        DEFAULT_BODY_PT,
                        base_dir=base_dir,
                        article_clean_html=article_clean_html,
                    )
            elif isinstance(ch, Tag):
                _emit_content_node(
                    doc,
                    ch,
                    base_dir,
                    list_level=list_level,
                    article_clean_html=article_clean_html,
                )
        return
    if "pdf-document-preview" in cls or "ppt-document-preview" in cls:
        for img in node.find_all("img"):
            if isinstance(img, Tag):
                if _should_skip_preview_image_for_docx(img):
                    continue
                _add_media_paragraph(
                    doc, img, base_dir, article_clean_html=article_clean_html
                )
        return
    if name in ("div", "section") and "word-preview-page" in cls:
        for ch in node.children:
            if isinstance(ch, NavigableString):
                if str(ch).strip():
                    bp = doc.add_paragraph()
                    _walk_inline_to_paragraph(
                        bp,
                        ch,
                        base_inline_ctx,
                        DEFAULT_BODY_PT,
                        base_dir=base_dir,
                        article_clean_html=article_clean_html,
                    )
            elif isinstance(ch, Tag):
                _emit_content_node(
                    doc,
                    ch,
                    base_dir,
                    list_level=list_level,
                    article_clean_html=article_clean_html,
                )
        return
    if name in _BLOCK_UNWRAP:
        for ch in node.children:
            if isinstance(ch, NavigableString):
                if str(ch).strip():
                    bp = doc.add_paragraph()
                    _walk_inline_to_paragraph(
                        bp,
                        ch,
                        base_inline_ctx,
                        DEFAULT_BODY_PT,
                        base_dir=base_dir,
                        article_clean_html=article_clean_html,
                    )
            elif isinstance(ch, Tag):
                _emit_content_node(
                    doc,
                    ch,
                    base_dir,
                    list_level=list_level,
                    article_clean_html=article_clean_html,
                )
        return
    if _is_block_media_tag(node):
        _add_media_paragraph(
            doc, node, base_dir, article_clean_html=article_clean_html
        )
        return
    if name in ("h1", "h2", "h3", "h4", "h5", "h6"):
        lvl = int(name[1])
        style = f"Heading {min(lvl, 9)}"
        hp = doc.add_paragraph(style=style)
        for ch in node.children:
            _walk_inline_to_paragraph(
                hp,
                ch,
                base_inline_ctx,
                DEFAULT_BODY_PT,
                base_dir=base_dir,
                article_clean_html=article_clean_html,
            )
        return
    if name == "ul":
        _emit_ul(
            doc,
            node,
            base_dir,
            list_level=list_level,
            article_clean_html=article_clean_html,
        )
        return
    if name == "ol":
        _emit_ol(
            doc,
            node,
            base_dir,
            list_level=list_level,
            article_clean_html=article_clean_html,
        )
        return
    if name == "table":
        _emit_html_table(
            doc, node, base_dir, article_clean_html=article_clean_html
        )
        return
    if name == "hr":
        hp = doc.add_paragraph()
        _apply_body_paragraph_format(hp)
        r = hp.add_run("―" * 28)
        _set_run_east_asia(r, "绛夌嚎")
        r.font.size = DEFAULT_BODY_PT
        r.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
        return
    if name == "blockquote":
        for ch in node.children:
            if isinstance(ch, NavigableString):
                if str(ch).strip():
                    bp = doc.add_paragraph()
                    bp.paragraph_format.left_indent = Inches(0.28)
                    _walk_inline_to_paragraph(
                        bp,
                        ch,
                        base_inline_ctx,
                        DEFAULT_BODY_PT,
                        base_dir=base_dir,
                        article_clean_html=article_clean_html,
                    )
            elif isinstance(ch, Tag):
                cn = (ch.name or "").lower()
                if cn in ("p", "div", "section"):
                    if cn in ("div", "section") and "word-preview-page" in _tag_class_list(
                        ch
                    ):
                        for im in ch.find_all("img"):
                            if isinstance(im, Tag):
                                _add_media_paragraph(
                                    doc,
                                    im,
                                    base_dir,
                                    article_clean_html=article_clean_html,
                                )
                        continue
                    bp = doc.add_paragraph()
                    bp.paragraph_format.left_indent = Inches(0.28)
                    for c2 in ch.children:
                        _walk_inline_to_paragraph(
                            bp,
                            c2,
                            base_inline_ctx,
                            DEFAULT_BODY_PT,
                            base_dir=base_dir,
                            article_clean_html=article_clean_html,
                        )
                else:
                    bp = doc.add_paragraph()
                    bp.paragraph_format.left_indent = Inches(0.28)
                    _walk_inline_to_paragraph(
                        bp,
                        ch,
                        base_inline_ctx,
                        DEFAULT_BODY_PT,
                        base_dir=base_dir,
                        article_clean_html=article_clean_html,
                    )
        return
    if name == "figure":
        for ch in node.children:
            if isinstance(ch, Tag) and (ch.name or "").lower() != "figcaption":
                _emit_content_node(
                    doc,
                    ch,
                    base_dir,
                    list_level=list_level,
                    article_clean_html=article_clean_html,
                )
        for ch in node.children:
            if isinstance(ch, Tag) and (ch.name or "").lower() == "figcaption":
                fp = doc.add_paragraph()
                fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for c2 in ch.children:
                    _walk_inline_to_paragraph(
                        fp,
                        c2,
                        base_inline_ctx,
                        META_PT,
                        base_dir=base_dir,
                        article_clean_html=article_clean_html,
                    )
        return
    if name in ("p", "div", "section"):
        has_block_child = any(
            isinstance(ch, Tag)
            and (
                (ch.name or "").lower()
                in ("div", "p", "section", "ul", "ol", "table", "blockquote", "figure")
            )
            for ch in node.children
        )
        if has_block_child:
            for ch in node.children:
                if isinstance(ch, NavigableString):
                    if str(ch).strip():
                        bp = doc.add_paragraph()
                        _apply_node_paragraph_css(bp, node)
                        _walk_inline_to_paragraph(
                            bp,
                            ch,
                            base_inline_ctx,
                            DEFAULT_BODY_PT,
                            base_dir=base_dir,
                            article_clean_html=article_clean_html,
                        )
                    continue
                if isinstance(ch, Tag):
                    _emit_content_node(
                        doc,
                        ch,
                        base_dir,
                        list_level=list_level,
                        article_clean_html=article_clean_html,
                    )
            return
        if node.find(["img", "video", "audio"]):
            seen_media_keys: set[str] = set()
            for sub in node.children:
                if isinstance(sub, Tag) and _is_block_media_tag(sub):
                    mk = _media_rel_key_for_dedupe(sub)
                    if mk is not None and mk in seen_media_keys:
                        continue
                    if mk is not None:
                        seen_media_keys.add(mk)
                    _add_media_paragraph(
                        doc,
                        sub,
                        base_dir,
                        article_clean_html=article_clean_html,
                    )
                elif isinstance(sub, Tag):
                    sn = (sub.name or "").lower()
                    if sn in ("div", "section") and "word-preview-page" in _tag_class_list(
                        sub
                    ):
                        for im in sub.find_all("img"):
                            if isinstance(im, Tag):
                                _add_media_paragraph(
                                    doc,
                                    im,
                                    base_dir,
                                    article_clean_html=article_clean_html,
                                )
                        continue
                    bp = doc.add_paragraph()
                    _apply_node_paragraph_css(bp, sub if isinstance(sub, Tag) else node)
                    _walk_inline_to_paragraph(
                        bp,
                        sub,
                        base_inline_ctx,
                        DEFAULT_BODY_PT,
                        base_dir=base_dir,
                        article_clean_html=article_clean_html,
                        local_media_href_dedupe=set(),
                    )
                elif isinstance(sub, NavigableString) and str(sub).strip():
                    bp = doc.add_paragraph()
                    _apply_node_paragraph_css(bp, node)
                    _walk_inline_to_paragraph(
                        bp,
                        sub,
                        base_inline_ctx,
                        DEFAULT_BODY_PT,
                        base_dir=base_dir,
                        article_clean_html=article_clean_html,
                    )
            return
        bp = doc.add_paragraph()
        _apply_node_paragraph_css(bp, node)
        for ch in node.children:
            _walk_inline_to_paragraph(
                bp,
                ch,
                base_inline_ctx,
                DEFAULT_BODY_PT,
                base_dir=base_dir,
                article_clean_html=article_clean_html,
                local_media_href_dedupe=set(),
            )
        return
    if name == "li":
        bp = doc.add_paragraph(style="List Bullet")
        ind = _list_indent_for_level(list_level)
        if ind is not None:
            bp.paragraph_format.left_indent = ind
        for ch in node.children:
            _walk_inline_to_paragraph(
                bp,
                ch,
                base_inline_ctx,
                DEFAULT_BODY_PT,
                base_dir=base_dir,
                article_clean_html=article_clean_html,
            )
        return
    bp = doc.add_paragraph()
    for ch in node.children:
        _walk_inline_to_paragraph(
            bp,
            ch,
            base_inline_ctx,
            DEFAULT_BODY_PT,
            base_dir=base_dir,
            article_clean_html=article_clean_html,
        )


_CSS_COMMENT_RE = re.compile(r"/\*.*?\*/", re.S)
_CSS_RULE_RE = re.compile(r"([^{}]+)\{([^{}]*)\}", re.S)


def _read_local_css_for_docx(clean_html_path: Path, href: str) -> str:
    raw = (href or "").strip()
    if not raw or raw.startswith(("http://", "https://", "data:")):
        return ""
    base = clean_html_path.parent
    cand = (base / raw).resolve()
    if cand.is_file():
        return cand.read_text(encoding="utf-8", errors="ignore")
    bare = Path(raw.replace("\\", "/")).name
    if not bare:
        return ""
    rd = clean_html_path.with_suffix("")
    cand2 = (rd / bare).resolve()
    if cand2.is_file():
        return cand2.read_text(encoding="utf-8", errors="ignore")
    return ""


def _build_docx_class_ctx_map(soup: BeautifulSoup, clean_html_path: Path) -> dict[str, dict[str, str]]:
    chunks: list[str] = []
    for st in soup.find_all("style"):
        if isinstance(st, Tag):
            txt = st.get_text("\n", strip=False)
            if txt:
                chunks.append(txt)
    for lk in soup.find_all("link"):
        if not isinstance(lk, Tag):
            continue
        rel = " ".join(lk.get("rel") or []).lower()
        if "stylesheet" not in rel:
            continue
        href = str(lk.get("href") or "").strip()
        if not href:
            continue
        css_txt = _read_local_css_for_docx(clean_html_path, href)
        if css_txt:
            chunks.append(css_txt)

    out: dict[str, dict[str, str]] = {}
    for css in chunks:
        css2 = _CSS_COMMENT_RE.sub("", css or "")
        for m in _CSS_RULE_RE.finditer(css2):
            selector_group = (m.group(1) or "").strip()
            decl = _css_decl_map(m.group(2) or "")
            if not decl:
                continue
            filtered = {
                k: v
                for k, v in decl.items()
                if k
                in {
                    "color",
                    "font-size",
                    "font-family",
                    "font-style",
                    "font-weight",
                    "text-decoration",
                    "text-align",
                    "text-indent",
                    "margin-left",
                    "padding-left",
                    "width",
                }
            }
            if not filtered:
                continue
            for raw_sel in selector_group.split(","):
                sel = raw_sel.strip()
                if not sel.startswith("."):
                    continue
                # Keep simple single-class selectors only, e.g. ".fs1", ".fc0".
                if any(ch in sel for ch in " >+~:#["):
                    continue
                mm = re.fullmatch(r"\.([A-Za-z_][\w-]*)", sel)
                if not mm:
                    continue
                cls = mm.group(1)
                prev = out.get(cls, {})
                prev.update(filtered)
                out[cls] = prev
    return out


def convert_clean_html_file_to_docx(
    clean_html_path: Path,
    docx_path: Path,
    *,
    force: bool,
) -> bool:
    if docx_path.exists() and not force:
        return False
    text = clean_html_path.read_text(encoding="utf-8", errors="replace")
    soup = BeautifulSoup(text, "html.parser")
    _DOCX_TLS.class_ctx_map = _build_docx_class_ctx_map(soup, clean_html_path)
    try:
        title_el = soup.select_one("#title")
        if title_el:
            title = title_el.get_text(strip=True)
        elif soup.title:
            title = soup.title.get_text(" ", strip=True)
        else:
            title = "无标题"
        author_el = soup.select_one("#author")
        author = author_el.get_text(strip=True) if author_el else ""
        date_el = soup.select_one("#date")
        pub = date_el.get_text(strip=True) if date_el else ""
        content = soup.select_one("#content")
        if content is None:
            raise ValueError("清洗 HTML 缺少 #content")

        doc = Document()
        _init_doc_typography(doc)
        base_dir = clean_html_path.parent

        tp = doc.add_paragraph()
        tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        tr = tp.add_run(title)
        _set_run_east_asia(tr, "黑体")
        tr.font.size = TITLE_PT
        tr.bold = False

        mp = doc.add_paragraph()
        mp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        meta_line = "  ".join(x for x in (author, pub) if x)
        mr = mp.add_run(meta_line or " ")
        _set_run_east_asia(mr, "宋体")
        mr.font.size = META_PT

        for child in list(content.children):
            if isinstance(child, NavigableString):
                if str(child).strip():
                    bp = doc.add_paragraph()
                    _walk_inline_to_paragraph(
                        bp,
                        child,
                        _RunCtx(),
                        DEFAULT_BODY_PT,
                        base_dir=base_dir,
                        article_clean_html=clean_html_path,
                    )
                continue
            if isinstance(child, Tag):
                _emit_content_node(
                    doc,
                    child,
                    base_dir,
                    list_level=0,
                    article_clean_html=clean_html_path,
                )

        _apply_body_paragraph_format_to_all(doc)
        _remove_empty_paragraphs(doc)
        doc.save(str(docx_path))
        return True
    finally:
        _DOCX_TLS.class_ctx_map = {}


def docx_path_for_article_html(raw_html_path: Path) -> Path:
    return raw_html_path.with_suffix(".docx")


def clean_html_path_for_raw(raw_html_path: Path) -> Path:
    return raw_html_path.with_name(f"{CLEAN_HTML_PREFIX}{raw_html_path.name}")


def article_raw_and_clean_paths(path: Path) -> tuple[Path, Path]:
    # If path is orphan clean_ HTML, return (raw, clean); otherwise return (raw, clean_html_path_for_raw(raw)).
    name = path.name
    if name.lower().startswith(CLEAN_HTML_PREFIX.lower()):
        raw = path.with_name(name[len(CLEAN_HTML_PREFIX) :])
        return raw, path
    return path, clean_html_path_for_raw(path)


def res_dir_for_clean(clean_path: Path) -> Path:
    return clean_path.with_suffix("")


def _remove_clean_outputs(clean_path: Path) -> None:
    if clean_path.is_file():
        clean_path.unlink(missing_ok=True)
    rd = res_dir_for_clean(clean_path)
    if rd.is_dir():
        shutil.rmtree(rd, ignore_errors=True)


def _remove_article_sidecars(raw_path: Path) -> None:
    cc = clean_html_path_for_raw(raw_path)
    _remove_clean_outputs(cc)


def process_one_article(
    path: Path,
    session: requests.Session,
    *,
    force_clean: bool,
    remove_original: bool,
    r_clean_only: bool,
    gen_docx: bool,
    force_docx: bool,
) -> tuple[str, bool]:
    # Return value: (status, did_clean_write).
    # status: skipped | processed | failed | skipped_docx
    raw_path, clean_path = article_raw_and_clean_paths(path)
    docx_path = docx_path_for_article_html(raw_path)
    article_id = extract_article_id(raw_path)
    source_url = guess_source_url(raw_path)
    article_dir_name = raw_path.parent.name or "unknown"

    if r_clean_only:
        if not gen_docx:
            log_warn(f"--r-c 要求同时启用 -w 以生成 Word；跳过 {path.name}")
            return "failed", False
        if clean_path.is_file():
            try:
                if raw_path.is_file() and _is_preview_clean_html_file(clean_path):
                    st, saved = try_direct_download_preview_document_for_word_only(
                        raw_path, session, force=force_docx
                    )
                    if st in {"processed", "skipped"} and saved is not None:
                        log_info(f"[原文档] {saved}")
                        _remove_article_sidecars(raw_path)
                        if raw_path.is_file():
                            raw_path.unlink(missing_ok=True)
                        return "processed", False
                if convert_clean_html_file_to_docx(
                    clean_path, docx_path, force=force_docx
                ):
                    log_info(f"[docx] {docx_path}")
                    _remove_article_sidecars(raw_path)
                    if raw_path.is_file():
                        raw_path.unlink(missing_ok=True)
                    return "processed", False
                return "skipped", False
            except Exception as exc:
                log_warn(f"docx 失败 {clean_path}: {exc}")
                return "failed", False

        try:
            src_html = raw_path if raw_path.is_file() else clean_path
            if not src_html.is_file():
                log_warn(f"--r-c 无可用 HTML（缺 raw 与 clean）: {path.name}")
                return "failed", False
            text = _decode_html_bytes(src_html.read_bytes())
            soup = BeautifulSoup(text, "html.parser")
            title, author, publish_date = extract_article_meta(soup)
            article_title = title or raw_path.stem
            content = resolve_content_node(
                soup=soup,
                raw_html=text,
                session=session,
                source_url=source_url,
                article_id=article_id,
                article_title=article_title,
                article_dir_name=article_dir_name,
            )
            if content is None:
                raise ValueError("未找到正文容器")
            clean_soup = build_clean_soup(title, author, publish_date, content)
            with tempfile.TemporaryDirectory() as tmp:
                tdir = Path(tmp)
                tmp_clean = tdir / f"{CLEAN_HTML_PREFIX}article.html"
                rs = localize_resources(
                    clean_soup,
                    source_url,
                    tmp_clean,
                    session,
                    article_id=article_id,
                    article_title=article_title,
                    article_dir_name=article_dir_name,
                )
                if rs.failed_urls:
                    append_clean_article_error_line(
                        f"{article_id}\tarticle={article_title}\tdir={article_dir_name}"
                        f"\tresource_failed={len(rs.failed_urls)}"
                    )
                tmp_clean.write_text(str(clean_soup), encoding="utf-8")
                if rs.downloaded > 0:
                    time.sleep(random.uniform(*AFTER_ARTICLE_WITH_RESOURCES_SLEEP_SEC))
                if _is_preview_content_node(content):
                    saved = download_original_preview_document(
                        session,
                        article_id=article_id,
                        source_url=source_url,
                        output_base_path=raw_path.with_suffix(""),
                    )
                    if saved is not None:
                        log_info(f"[原文档] {saved}")
                        _remove_article_sidecars(raw_path)
                        if raw_path.is_file():
                            raw_path.unlink(missing_ok=True)
                        return "processed", False
                if convert_clean_html_file_to_docx(
                    tmp_clean, docx_path, force=True
                ):
                    log_info(f"[docx] {docx_path}")
                    _remove_article_sidecars(raw_path)
                    if raw_path.is_file():
                        raw_path.unlink(missing_ok=True)
                    return "processed", False
                return "failed", False
        except CleanRateLimitError:
            raise
        except Exception as exc:
            log_warn(f"清洗/转换失败 {path.name}: {exc}")
            append_clean_article_error_line(
                f"{article_id}\tarticle={raw_path.stem}\tdir={article_dir_name}"
                f"\tstatus=clean_failed\terr={exc}"
            )
            _remove_clean_outputs(clean_path)
            return "failed", False

    # Normal path: write clean_ prefixed HTML to disk.
    if not force_clean and clean_path.is_file():
        if gen_docx:
            try:
                if raw_path.is_file() and _is_preview_clean_html_file(clean_path):
                    st, saved = try_direct_download_preview_document_for_word_only(
                        raw_path, session, force=force_docx
                    )
                    if st in {"processed", "skipped"} and saved is not None:
                        log_info(f"[原文档] {saved}")
                        if remove_original and raw_path.is_file():
                            raw_path.unlink(missing_ok=True)
                        return "processed", False
                if convert_clean_html_file_to_docx(
                    clean_path, docx_path, force=force_docx
                ):
                    log_info(f"[docx] {docx_path}")
                    if remove_original and raw_path.is_file():
                        raw_path.unlink(missing_ok=True)
                    return "processed", False
            except Exception as exc:
                log_warn(f"docx 失败 {clean_path}: {exc}")
        return "skipped", False

    try:
        src_html = raw_path if raw_path.is_file() else clean_path
        if not src_html.is_file():
            log_warn(f"无可用 HTML: {path.name}")
            return "failed", False
        text = _decode_html_bytes(src_html.read_bytes())
        soup = BeautifulSoup(text, "html.parser")
        title, author, publish_date = extract_article_meta(soup)
        article_title = title or raw_path.stem
        content = resolve_content_node(
            soup=soup,
            raw_html=text,
            session=session,
            source_url=source_url,
            article_id=article_id,
            article_title=article_title,
            article_dir_name=article_dir_name,
        )
        if content is None:
            raise ValueError("未找到正文容器")
        clean_soup = build_clean_soup(title, author, publish_date, content)
        rs = localize_resources(
            clean_soup,
            source_url,
            clean_path,
            session,
            article_id=article_id,
            article_title=article_title,
            article_dir_name=article_dir_name,
        )
        if rs.failed_urls:
            append_clean_article_error_line(
                f"{article_id}\tarticle={article_title}\tdir={article_dir_name}"
                f"\tresource_failed={len(rs.failed_urls)}"
            )
        clean_path.write_text(str(clean_soup), encoding="utf-8")
        if rs.downloaded > 0:
            time.sleep(random.uniform(*AFTER_ARTICLE_WITH_RESOURCES_SLEEP_SEC))
        log_info(f"已清洗: {raw_path.name} -> {clean_path.name}")
        wrote = True
        if remove_original and raw_path.is_file():
            raw_path.unlink(missing_ok=True)
        if gen_docx:
            try:
                if _is_preview_content_node(content):
                    saved = download_original_preview_document(
                        session,
                        article_id=article_id,
                        source_url=source_url,
                        output_base_path=raw_path.with_suffix(""),
                    )
                    if saved is not None:
                        log_info(f"[原文档] {saved}")
                        return "processed", wrote
                if convert_clean_html_file_to_docx(
                    clean_path, docx_path, force=force_docx
                ):
                    log_info(f"[docx] {docx_path}")
            except Exception as exc:
                log_warn(f"docx 失败 {clean_path}: {exc}")
        return "processed", wrote
    except CleanRateLimitError:
        _remove_clean_outputs(clean_path)
        raise
    except Exception as exc:
        log_warn(f"清洗失败 {path.name}: {exc}")
        append_clean_article_error_line(
            f"{article_id}\tarticle={raw_path.stem}\tdir={article_dir_name}"
            f"\tstatus=clean_failed\terr={exc}"
        )
        _remove_clean_outputs(clean_path)
        return "failed", False


def docx_from_raw_html_via_temp(
    path: Path,
    session: requests.Session,
    *,
    force_docx: bool,
    remove_original: bool = False,
) -> str:
    # Word-only path: clean in temp dir without writing clean_ HTML to output tree.
    docx_path = docx_path_for_article_html(path)
    if docx_path.exists() and not force_docx:
        return "skipped"
    article_id = extract_article_id(path)
    source_url = guess_source_url(path)
    article_dir_name = path.parent.name or "unknown"
    try:
        text = _decode_html_bytes(path.read_bytes())
        soup = BeautifulSoup(text, "html.parser")
        title, author, publish_date = extract_article_meta(soup)
        article_title = title or path.stem
        content = resolve_content_node(
            soup=soup,
            raw_html=text,
            session=session,
            source_url=source_url,
            article_id=article_id,
            article_title=article_title,
            article_dir_name=article_dir_name,
        )
        if content is None:
            raise ValueError("未找到正文容器")
        clean_soup = build_clean_soup(title, author, publish_date, content)
        with tempfile.TemporaryDirectory() as tmp:
            tdir = Path(tmp)
            tmp_clean = tdir / f"{CLEAN_HTML_PREFIX}article.html"
            rs = localize_resources(
                clean_soup,
                source_url,
                tmp_clean,
                session,
                article_id=article_id,
                article_title=article_title,
                article_dir_name=article_dir_name,
            )
            if rs.failed_urls:
                append_clean_article_error_line(
                    f"{article_id}\tarticle={article_title}\tdir={article_dir_name}"
                    f"\tresource_failed={len(rs.failed_urls)}"
                )
            tmp_clean.write_text(str(clean_soup), encoding="utf-8")
            if rs.downloaded > 0:
                time.sleep(random.uniform(*AFTER_ARTICLE_WITH_RESOURCES_SLEEP_SEC))
            if _is_preview_content_node(content):
                saved = download_original_preview_document(
                    session,
                    article_id=article_id,
                    source_url=source_url,
                    output_base_path=path.with_suffix(""),
                )
                if saved is not None:
                    log_info(f"[原文档] {saved}")
                    if remove_original and path.is_file():
                        path.unlink(missing_ok=True)
                    return "processed"
            if convert_clean_html_file_to_docx(tmp_clean, docx_path, force=True):
                log_info(f"[docx] {docx_path}")
                if remove_original and path.is_file():
                    path.unlink(missing_ok=True)
                return "processed"
            return "failed"
    except CleanRateLimitError:
        raise
    except Exception as exc:
        log_warn(f"仅 Word 转换失败 {path}: {exc}")
        append_clean_article_error_line(
            f"{article_id}\tarticle={path.stem}\tdir={article_dir_name}"
            f"\tstatus=clean_failed\terr={exc}"
        )
        return "failed"


def run_clean_and_word_pass(
    root: Path,
    session: requests.Session,
    *,
    enable_clean: bool,
    gen_word: bool,
    force_clean: bool,
    force_docx: bool,
    remove_original: bool,
    r_clean_only: bool,
    limit: int = 0,
    remove_raw_when_word_only: bool = False,
    clean_article_pacing_sec: tuple[float, float] | None = None,
    offline_word_only: bool = False,
) -> int:
    if not enable_clean and not gen_word:
        return 0
    files = iter_library_article_html_files(root)
    if limit > 0:
        files = files[:limit]
    if not files:
        log_warn("未找到待处理的文库 HTML。")
        return 0
    ok = fail = skip = 0
    use_clean_pacing = bool(clean_article_pacing_sec and (enable_clean or r_clean_only))
    pacing_lo = pacing_hi = 0.0
    if use_clean_pacing:
        assert clean_article_pacing_sec is not None
        pacing_lo, pacing_hi = clean_article_pacing_sec
        if pacing_lo < 0:
            pacing_lo = 0.0
        if pacing_hi < pacing_lo:
            pacing_hi = pacing_lo

    # Local word-only mode: convert existing clean_ HTML directly with max workers.
    if gen_word and not enable_clean and not r_clean_only:
        max_workers = max(1, RESOURCE_DOWNLOAD_MAX_WORKERS)
        tasks: list[tuple[str, Path, Path, Path]] = []
        for fp in files:
            raw_fp, clean_fp = article_raw_and_clean_paths(fp)
            docx_path = docx_path_for_article_html(raw_fp)
            if not clean_fp.is_file():
                if raw_fp.is_file() and not offline_word_only:
                    tasks.append(("preview_download", raw_fp, clean_fp, docx_path))
                else:
                    log_warn(f"本地 Word 模式跳过（缺 clean 文件）: {raw_fp.name}")
                    skip += 1
                continue
            tasks.append(("clean_to_docx", raw_fp, clean_fp, docx_path))

        def _word_local_worker(
            item: tuple[str, Path, Path, Path]
        ) -> tuple[str, Path, Path | None]:
            mode, raw_fp, clean_fp, docx_path = item
            try:
                if mode == "clean_to_docx":
                    if (
                        (not offline_word_only)
                        and raw_fp.is_file()
                        and _is_preview_clean_html_file(clean_fp)
                    ):
                        st, saved = try_direct_download_preview_document_for_word_only(
                            raw_fp, session, force=force_docx
                        )
                        if st == "processed":
                            return "processed_download", raw_fp, saved
                        if st == "skipped":
                            return "skipped", raw_fp, saved
                    if convert_clean_html_file_to_docx(
                        clean_fp, docx_path, force=force_docx
                    ):
                        return "processed", raw_fp, docx_path
                    return "skipped", raw_fp, docx_path

                # Fallback for preview-only raw pages: download original doc/pdf/ppt directly.
                st, saved = try_direct_download_preview_document_for_word_only(
                    raw_fp, session, force=force_docx
                )
                if st == "processed":
                    return "processed_download", raw_fp, saved
                if st == "skipped":
                    return "skipped", raw_fp, saved
                if st == "not_preview":
                    log_warn(f"本地 Word 模式跳过（缺 clean 且非预览文档）: {raw_fp.name}")
                    return "skipped", raw_fp, None
                return "failed", raw_fp, None
            except Exception as exc:
                log_warn(f"Word 本地处理失败 {raw_fp.name}: {exc}")
                return "failed", raw_fp, None

        with concurrent.futures.ThreadPoolExecutor(max_workers=max_workers) as ex:
            futures = [ex.submit(_word_local_worker, item) for item in tasks]
            for idx, fut in enumerate(concurrent.futures.as_completed(futures), start=1):
                st, raw_fp, out_path = fut.result()
                if st == "processed":
                    if out_path is not None:
                        log_info(f"[docx] {out_path}")
                    if remove_raw_when_word_only and raw_fp.is_file():
                        raw_fp.unlink(missing_ok=True)
                    ok += 1
                elif st == "processed_download":
                    if out_path is not None:
                        log_info(f"[原文档] {out_path}")
                    if remove_raw_when_word_only and raw_fp.is_file():
                        raw_fp.unlink(missing_ok=True)
                    ok += 1
                elif st == "skipped":
                    skip += 1
                else:
                    fail += 1
                if idx % 20 == 0:
                    log_info(f"清洗/Word 进度: {idx}/{len(tasks)}")

        log_info(
            f"清洗/Word 完成: 处理 {ok} / 跳过 {skip} / 失败 {fail} / 共 {len(files)}"
        )
        return fail

    for idx, fp in enumerate(files, start=1):
        should_pace = use_clean_pacing
        if r_clean_only and gen_word:
            _, clean_fp = article_raw_and_clean_paths(fp)
            local_clean_ready = clean_fp.is_file() and (not force_clean)
            st, _ = process_one_article(
                fp,
                session,
                force_clean=force_clean,
                remove_original=False,
                r_clean_only=True,
                gen_docx=True,
                force_docx=force_docx,
            )
            if local_clean_ready and st in {"processed", "skipped"}:
                should_pace = False
            if st == "processed":
                ok += 1
            elif st == "skipped":
                skip += 1
            else:
                fail += 1
        elif enable_clean:
            _, clean_fp = article_raw_and_clean_paths(fp)
            local_clean_ready = clean_fp.is_file() and (not force_clean)
            st, _ = process_one_article(
                fp,
                session,
                force_clean=force_clean,
                remove_original=remove_original,
                r_clean_only=r_clean_only,
                gen_docx=gen_word,
                force_docx=force_docx,
            )
            if local_clean_ready and st in {"processed", "skipped"}:
                should_pace = False
            if st == "processed":
                ok += 1
            elif st == "skipped":
                skip += 1
            else:
                fail += 1
        elif gen_word:
            raw_fp, clean_fp = article_raw_and_clean_paths(fp)
            docx_path = docx_path_for_article_html(raw_fp)
            if clean_fp.is_file():
                try:
                    if raw_fp.is_file() and _is_preview_clean_html_file(clean_fp):
                        st2, saved = try_direct_download_preview_document_for_word_only(
                            raw_fp, session, force=force_docx
                        )
                        if st2 == "processed":
                            if saved is not None:
                                log_info(f"[原文档] {saved}")
                            if remove_raw_when_word_only and raw_fp.is_file():
                                raw_fp.unlink(missing_ok=True)
                            ok += 1
                            continue
                        if st2 == "skipped":
                            skip += 1
                            continue
                    if convert_clean_html_file_to_docx(
                        clean_fp, docx_path, force=force_docx
                    ):
                        log_info(f"[docx] {docx_path}")
                        if remove_raw_when_word_only and raw_fp.is_file():
                            raw_fp.unlink(missing_ok=True)
                        ok += 1
                    else:
                        skip += 1
                except Exception as exc:
                    log_warn(f"docx 失败 {clean_fp}: {exc}")
                    fail += 1
            elif raw_fp.is_file():
                st = docx_from_raw_html_via_temp(
                    raw_fp,
                    session,
                    force_docx=force_docx,
                    remove_original=remove_raw_when_word_only,
                )
                if st == "processed":
                    ok += 1
                elif st == "skipped":
                    skip += 1
                else:
                    fail += 1
            else:
                log_warn(f"跳过（无 raw 与 clean HTML）: {fp}")
                fail += 1
        if idx % 20 == 0:
            log_info(f"清洗/Word 进度: {idx}/{len(files)}")
        if should_pace and idx < len(files):
            time.sleep(random.uniform(pacing_lo, pacing_hi))
    log_info(
        f"清洗/Word 完成: 处理 {ok} / 跳过 {skip} / 失败 {fail} / 共 {len(files)}"
    )
    return fail


def _extract_not_found_log_entry(line: str) -> tuple[str, str] | None:
    m_id = re.search(r"article_id=([0-9]+)", line)
    m_res = re.search(r"resource=([^\t]+)", line)
    if not m_id or not m_res:
        return None
    return m_id.group(1).strip(), m_res.group(1).strip()


def _extract_clean_error_log_entry(line: str) -> tuple[str, str] | None:
    prefix = line.split("\t", 1)[0]
    m = re.match(r"^([0-9]+)-(https?://.+?)-", prefix)
    if not m:
        return None
    return m.group(1).strip(), m.group(2).strip()


def _find_article_html_by_id(root: Path, article_id: str) -> Path | None:
    raw_files = sorted(root.rglob(f"{article_id}-*.html"))
    for p in raw_files:
        if p.is_file() and not p.name.lower().startswith(CLEAN_HTML_PREFIX.lower()):
            return p
    clean_files = sorted(root.rglob(f"{CLEAN_HTML_PREFIX}{article_id}-*.html"))
    for p in clean_files:
        if p.is_file():
            return p
    return None


def _probe_resource_recoverable(
    session: requests.Session,
    article_id: str,
    resource_url: str,
) -> bool:
    source_url = f"{BASE_URL}/showweb/0/0/{article_id}.aspx"

    def _try_fetch(one_url: str) -> bool:
        if not one_url:
            return False
        try:
            request_with_retry(
                session,
                one_url,
                headers={
                    "Referer": source_url,
                    "Accept": "image/avif,image/webp,image/apng,image/*,*/*;q=0.8",
                    "Accept-Language": "zh-CN,zh;q=0.9,en;q=0.8",
                    "Sec-Fetch-Dest": "image",
                    "Sec-Fetch-Mode": "no-cors",
                    "Sec-Fetch-Site": "cross-site",
                    "User-Agent": session.headers.get("User-Agent", "Mozilla/5.0"),
                },
                timeout=RESOURCE_REQUEST_TIMEOUT,
                retries=RESOURCE_REQUEST_RETRIES,
                use_session_cookies=("signature=" not in one_url.lower()),
                bypass_env_proxy=True,
            )
            return True
        except Exception:
            return False

    if _try_fetch(resource_url):
        return True

    try:
        by_input, by_path = _request_changeurl_signed_images(
            session, source_url, [resource_url]
        )
        fresh = by_input.get(resource_url, "")
        if not fresh:
            fresh = by_path.get(_url_path_key(resource_url), "")
        if fresh and _try_fetch(fresh):
            return True
    except Exception:
        pass
    return False


def replay_resource_failures_from_logs(
    root: Path,
    session: requests.Session,
) -> dict[str, int]:
    stats = {
        "entries_total": 0,
        "entries_recoverable": 0,
        "articles_retried": 0,
        "articles_recleaned": 0,
        "lines_removed": 0,
    }

    log_rows: list[tuple[str, int, str, str, str, str]] = []
    # (file_key, line_idx, raw_line, parser_kind, article_id, resource_url)
    file_map = {
        "not_found": RESOURCES_NOT_FOUND_WARNING_FILE,
        "clean_error": CLEAN_ERROR_URL_FILE,
    }
    parser_map = {
        "not_found": _extract_not_found_log_entry,
        "clean_error": _extract_clean_error_log_entry,
    }

    for key, fp in file_map.items():
        if not fp.is_file():
            continue
        lines = fp.read_text(encoding="utf-8", errors="ignore").splitlines()
        parser = parser_map[key]
        for idx, line in enumerate(lines):
            parsed = parser(line)
            if not parsed:
                continue
            aid, res = parsed
            if not aid or not res:
                continue
            log_rows.append((key, idx, line, key, aid, res))

    stats["entries_total"] = len(log_rows)
    if not log_rows:
        return stats

    pair_probe_cache: dict[tuple[str, str], bool] = {}
    for _, _, _, _, aid, res in log_rows:
        pair = (aid, res)
        if pair in pair_probe_cache:
            continue
        pair_probe_cache[pair] = _probe_resource_recoverable(session, aid, res)

    recoverable_pairs = {p for p, ok in pair_probe_cache.items() if ok}
    stats["entries_recoverable"] = len(recoverable_pairs)
    if not recoverable_pairs:
        return stats

    recoverable_articles = sorted({aid for aid, _ in recoverable_pairs})
    recleaned_articles: set[str] = set()
    for aid in recoverable_articles:
        fp = _find_article_html_by_id(root, aid)
        if fp is None:
            continue
        stats["articles_retried"] += 1
        try:
            st, _ = process_one_article(
                fp,
                session,
                force_clean=True,
                remove_original=False,
                r_clean_only=False,
                gen_docx=False,
                force_docx=False,
            )
        except CleanRateLimitError:
            raise
        except Exception as exc:
            log_warn(f"日志回放复洗失败 art={aid} file={fp.name} err={exc}")
            continue
        if st == "processed":
            recleaned_articles.add(aid)

    stats["articles_recleaned"] = len(recleaned_articles)
    if not recleaned_articles:
        return stats

    to_remove: dict[str, set[int]] = {"not_found": set(), "clean_error": set()}
    for key, idx, _, _, aid, res in log_rows:
        if aid in recleaned_articles and (aid, res) in recoverable_pairs:
            to_remove[key].add(idx)

    for key, fp in file_map.items():
        if not fp.is_file():
            continue
        rm = to_remove.get(key) or set()
        if not rm:
            continue
        lines = fp.read_text(encoding="utf-8", errors="ignore").splitlines()
        kept = [ln for i, ln in enumerate(lines) if i not in rm]
        fp.write_text(("\n".join(kept) + ("\n" if kept else "")), encoding="utf-8")
        stats["lines_removed"] += len(rm)

    return stats


if __name__ == "__main__":
    raise SystemExit("入口：仓库根 wc-library.py 或 src/wc-library.py（本文件为库模块，不作为主程序）。")

